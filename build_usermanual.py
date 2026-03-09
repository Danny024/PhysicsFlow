"""
PhysicsFlow User Manual Generator — python-docx.

Run:  python build_usermanual.py
Output: PhysicsFlow_UserManual.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0A, 0x16, 0x28)
BLUE      = RGBColor(0x15, 0x65, 0xC0)
CYAN      = RGBColor(0x00, 0xB0, 0xD8)
GREEN     = RGBColor(0x1B, 0x8A, 0x3C)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY  = RGBColor(0x66, 0x66, 0x66)
RED       = RGBColor(0xC0, 0x22, 0x22)
AMBER     = RGBColor(0xE6, 0x8A, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_GREY


# ── Utility helpers ────────────────────────────────────────────────────────────

def h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = color
    _add_bottom_border(p, '1565C0')
    return p


def h2(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def h3(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def body(text, italic=False, color=DARK_GREY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def bullet(text, level=1, color=DARK_GREY):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 * level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = color
    return p


def numbered(text, level=1, color=DARK_GREY):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 * level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = color
    return p


def note(text, color=BLUE):
    """Callout box styled as an indented note."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run('📝  NOTE:  ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = color
    run2 = p.add_run(text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = DARK_GREY
    return p


def warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run('⚠  WARNING:  ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RED
    run2 = p.add_run(text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = DARK_GREY
    return p


def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run('💡  TIP:  ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = GREEN
    run2 = p.add_run(text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = DARK_GREY
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x3A)
    return p


def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                cell.width = col_widths[j]
    # Header row
    for j, h in enumerate(headers):
        c = tbl.cell(0, j)
        _shade_cell(c, '0A1628')
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = CYAN
        run.font.name = 'Calibri'
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.cell(i + 1, j)
            _shade_cell(c, 'F8FAFC' if i % 2 == 0 else 'FFFFFF')
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GREY
            run.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_bottom_border(para, hex_color='1565C0'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(60)
p_title.paragraph_format.space_after  = Pt(6)
r = p_title.add_run('PhysicsFlow')
r.font.name = 'Calibri'
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
r2 = p_sub.add_run('AI-Native Reservoir Simulation & History Matching')
r2.font.name = 'Calibri'
r2.font.size = Pt(16)
r2.font.color.rgb = BLUE

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_after = Pt(0)
r3 = p_ver.add_run('User Manual  —  Version 1.3.0')
r3.font.name = 'Calibri'
r3.font.size = Pt(13)
r3.font.color.rgb = MID_GREY

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_after = Pt(40)
r4 = p_date.add_run('March 2026')
r4.font.name = 'Calibri'
r4.font.size = Pt(11)
r4.font.color.rgb = MID_GREY
r4.font.italic = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

h1('Table of Contents')
toc_items = [
    ('1', 'Introduction', '3'),
    ('2', 'System Requirements', '4'),
    ('3', 'Installation', '5'),
    ('4', 'Launching PhysicsFlow', '7'),
    ('5', 'Creating a New Project', '8'),
    ('6', 'Training the PINO Surrogate', '11'),
    ('7', 'History Matching (αREKI)', '13'),
    ('8', 'Production Forecast', '15'),
    ('9', '3D Reservoir Viewer', '17'),
    ('10', '2D Cross-Section Viewer', '19'),
    ('11', 'Reports & Excel Export', '20'),
    ('12', 'AI Reservoir Assistant', '22'),
    ('13', 'Knowledge Base Management (RAG + KG)', '25'),
    ('14', 'Project File Encryption', '27'),
    ('15', 'PINO Pre-Training CLI', '29'),
    ('16', 'Settings & Configuration', '30'),
    ('17', 'Troubleshooting', '31'),
    ('18', 'Keyboard Shortcuts', '33'),
    ('19', 'Glossary', '34'),
]
toc_tbl = doc.add_table(rows=len(toc_items), cols=3)
toc_tbl.style = 'Table Grid'
for i, (num, title, page) in enumerate(toc_items):
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for j in range(3):
        _shade_cell(toc_tbl.cell(i, j), bg)
    toc_tbl.cell(i, 0).width = Cm(1.5)
    toc_tbl.cell(i, 1).width = Cm(14.0)
    toc_tbl.cell(i, 2).width = Cm(1.5)
    def _tp(cell, text, bold=False, right=False, col=DARK_GREY):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.color.rgb = col
    _tp(toc_tbl.cell(i, 0), num, bold=True, col=BLUE)
    _tp(toc_tbl.cell(i, 1), title)
    _tp(toc_tbl.cell(i, 2), page, right=True, col=MID_GREY)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1('1  Introduction')
body('PhysicsFlow is an AI-accelerated reservoir simulation and history matching platform. '
     'It replaces classical finite-volume simulators in the history matching loop with a '
     'Physics-Informed Neural Operator (PINO) surrogate, achieving a ~6,000× speed-up '
     'while maintaining physical consistency via Darcy PDE loss.')

body('This manual covers everything required to install, configure, and use PhysicsFlow v1.3.0 '
     'from project creation through history matching, forecasting, 3D visualisation, and the '
     'v1.3.0 Intelligence Layer — Hybrid RAG knowledge assistant and Reservoir Knowledge Graph.')

h2('1.1  Key Capabilities')
simple_table(
    ['Feature', 'Description'],
    [
        ('PINO Forward Simulation', '~7 seconds per full-field run vs. 3–8 hours in Eclipse 100'),
        ('αREKI History Matching', 'Adaptive Ensemble Kalman Inversion — 200 members, ~1 hour'),
        ('Uncertainty Quantification', 'P10/P50/P90 fan charts, VCAE + DDIM generative priors'),
        ('3D Reservoir Viewer', 'Interactive HelixToolkit voxel renderer with animation'),
        ('2D Cross-Section Viewer', 'I/J/K-plane slices with Jet/Viridis/Seismic/Greys colormaps'),
        ('AI Reservoir Assistant', 'Local LLM (Ollama) with 10 tools: 8 live + RAG document search + KG query'),
        ('Hybrid RAG', 'ChromaDB + BM25 + cross-encoder reranking over indexed project documents'),
        ('Reservoir Knowledge Graph', 'Structured Norne topology: 22 wells, 22 layers, 53 faults — 20-pattern NL queries'),
        ('Reports & Export', 'QuestPDF HM/EUR reports; ClosedXML Excel with well/ensemble/training data'),
        ('Project File Encryption', 'AES-256-GCM (.pfproj.enc) with PBKDF2-HMAC-SHA256'),
        ('Eclipse I/O', 'Native .DATA / .EGRID / .UNRST reader; LAS 2.0 well logs'),
        ('Database', 'SQLite audit trail: projects, runs, epochs, iterations, observations'),
    ],
    col_widths=[Cm(5.5), Cm(12.0)],
)

h2('1.2  Supported Platforms')
body('PhysicsFlow v1.3.0 runs on:')
bullet('Windows 10 (build 19041+) or Windows 11, 64-bit')
bullet('NVIDIA GPU with CUDA 12.x — strongly recommended for training and ensemble operations')
bullet('CPU-only mode is supported but training will be significantly slower')
note('macOS and Linux are not supported in the desktop application. The Python engine '
     '(physicsflow) can be installed on Linux for headless pre-training and CLI use.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

h1('2  System Requirements')

h2('2.1  Minimum Requirements')
simple_table(
    ['Component', 'Minimum', 'Recommended'],
    [
        ('OS', 'Windows 10 64-bit (build 19041)', 'Windows 11 64-bit'),
        ('CPU', 'Intel Core i5 / AMD Ryzen 5 (8th gen+)', 'Intel Core i7/i9 or Ryzen 7/9'),
        ('RAM', '16 GB', '32 GB or more'),
        ('GPU', 'None (CPU fallback)', 'NVIDIA RTX 3080 or better (12 GB+ VRAM)'),
        ('Storage', '20 GB free space', '50 GB SSD'),
        ('Display', '1920 × 1080', '2560 × 1440 or higher'),
        ('.NET Runtime', '.NET 8.0 Desktop Runtime', 'Bundled in installer — no manual install'),
        ('Python', 'Python 3.11 (bundled)', 'Bundled — no separate installation required'),
        ('Ollama', 'Optional — for AI assistant', 'Ollama 0.2+ with phi3:mini pulled'),
    ],
    col_widths=[Cm(4.0), Cm(6.0), Cm(7.5)],
)

h2('2.2  CUDA Requirements')
body('GPU acceleration requires NVIDIA CUDA 12.x. The PhysicsFlow installer bundles pre-compiled '
     'PyTorch CUDA wheels — no separate CUDA toolkit installation is required for the desktop '
     'application. For the developer setup, CUDA 12.x toolkit is needed.')

note('NVIDIA driver version 525.85.12 or later is required for CUDA 12.x support. '
     'Check your driver version in NVIDIA Control Panel or with: nvidia-smi')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════

h1('3  Installation')

h2('3.1  Production Installer (Recommended)')
body('The self-contained WiX bootstrapper bundles all dependencies. No manual Python or .NET '
     'installation is required for most users.')
numbered('Download PhysicsFlow-Installer-1.3.0-x64.exe from the releases page.')
numbered('Double-click the installer. If prompted by Windows SmartScreen, click "More info" then "Run anyway".')
numbered('The bootstrapper automatically installs the following prerequisites if missing:')
bullet('.NET 8 Desktop Runtime', level=2)
bullet('Visual C++ 2022 Redistributable (x64)', level=2)
numbered('Choose your installation directory (default: C:\\Program Files\\PhysicsFlow).')
numbered('Click Install. The installer will:')
bullet('Copy the desktop application and bundled Python environment', level=2)
bullet('Register the .pfproj file association', level=2)
bullet('Create a Start Menu shortcut', level=2)
bullet('Run pip install for the physicsflow Python package', level=2)
numbered('Click Finish. PhysicsFlow is ready to launch.')

tip('The installer requires approximately 3 GB of disk space including PyTorch wheels. '
    'Ensure you have at least 5 GB free before starting.')

h2('3.2  Installing the AI Assistant (Ollama)')
body('The AI Reservoir Assistant requires Ollama to be installed and running on your machine. '
     'Your data never leaves your computer — all inference is local.')
numbered('Download Ollama from ollama.com and run the installer.')
numbered('Open a terminal and pull the recommended model:')
code_block('ollama pull phi3:mini')
numbered('Verify Ollama is running:')
code_block('ollama list')
numbered('In PhysicsFlow, go to Settings → AI Assistant and confirm the Ollama endpoint '
         '(default: http://localhost:11434).')

note('A larger model like llama3:8b will give better responses but requires 8+ GB VRAM. '
     'phi3:mini runs well on CPU with 8 GB RAM.')

h2('3.3  Developer Setup')
body('For developers who want to modify the source code or run the Python engine independently:')
numbered('Clone the repository:')
code_block('git clone https://github.com/Danny024/PhysicsFlow.git\ncd PhysicsFlow')
numbered('Create and activate the Python virtual environment:')
code_block('cd engine\npython -m venv .venv\n.venv\\Scripts\\activate')
numbered('Install the package in editable mode:')
code_block('pip install -e ".[dev]"')
numbered('Generate gRPC stubs from Protocol Buffer definitions:')
code_block('cd physicsflow/proto\npython -m grpc_tools.protoc -I. --python_out=. --grpc_python_out=. *.proto')
numbered('Build and run the .NET application:')
code_block('cd ../../desktop\ndotnet restore\ndotnet run --project src/PhysicsFlow.App')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. LAUNCHING PHYSICSFLOW
# ══════════════════════════════════════════════════════════════════════════════

h1('4  Launching PhysicsFlow')

h2('4.1  Starting the Application')
body('Launch PhysicsFlow from the Start Menu or by double-clicking a .pfproj project file. '
     'On first launch, the application will:')
bullet('Start the Python gRPC engine in the background')
bullet('Wait for the engine.ready signal file (typically 3–8 seconds)')
bullet('Display the Dashboard once the engine is connected')

body('The status bar at the bottom shows the engine connection status. A green indicator '
     'means the engine is running and ready.')

h2('4.2  The Main Interface')
body('PhysicsFlow uses a three-column layout:')

simple_table(
    ['Panel', 'Location', 'Description'],
    [
        ('Navigation Sidebar', 'Left column', 'Click to switch between Dashboard, Project Setup, Training, History Matching, Forecast, Visualisation'),
        ('Main Content Area', 'Centre column', 'The active view (changes with navigation)'),
        ('AI Assistant Panel', 'Right column', 'Always-visible chat panel for asking questions about your reservoir'),
    ],
    col_widths=[Cm(4.0), Cm(4.0), Cm(9.5)],
)

h2('4.3  Navigation')
body('Click any item in the left sidebar to navigate:')
bullet('Dashboard — project statistics, recent well performance, quick actions')
bullet('Project Setup — 5-step project wizard (grid, wells, PVT, schedule, save)')
bullet('Training — PINO surrogate training monitor with live loss curves')
bullet('History Matching — αREKI workspace, ensemble fan chart, per-well heatmap')
bullet('Forecast — P10/P50/P90 fan charts, EUR statistics, PDF and Excel export')
bullet('3D Viewer — interactive 3D voxel renderer for the reservoir model')
bullet('Cross-Section — 2D I/J/K-plane slice viewer')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. CREATING A NEW PROJECT
# ══════════════════════════════════════════════════════════════════════════════

h1('5  Creating a New Project')

body('Every study in PhysicsFlow is captured in a project file (.pfproj). '
     'The 5-step wizard guides you through the setup.')

h2('5.1  Starting the Wizard')
numbered('Click New Project on the Dashboard or select File → New Project.')
numbered('The Project Setup wizard opens on Step 1: Grid.')

h2('5.2  Step 1 — Grid')
body('Define the reservoir grid dimensions:')
bullet('Enter Nx, Ny, Nz (number of cells in each direction)')
bullet('Or click Import Eclipse Deck and select a .DATA file — PhysicsFlow will read the '
       'DIMENS keyword and populate all grid properties automatically')
bullet('Review the active cell count displayed below the grid inputs')

note('The Norne benchmark field uses a 46 × 112 × 22 grid (113,344 cells total, '
     '~44,000 active after fault removal).')

h2('5.3  Step 2 — Wells')
body('Define producer and injector wells:')
bullet('Click Import COMPDAT to read well connections from an Eclipse .DATA file')
bullet('Or click Load Norne Defaults to populate the 22 producers, 9 water injectors, '
       'and 4 gas injectors from the Norne benchmark')
bullet('Manually add wells by clicking + Add Well and entering the well name, type, '
       'and I/J/K perforation intervals')
bullet('Click any well to edit its COMPDAT connections and production constraints')

h2('5.4  Step 3 — PVT')
body('Configure the Black-Oil PVT properties:')

simple_table(
    ['Property', 'Description', 'Norne Default'],
    [
        ('Initial Pressure (Pi)', 'Reservoir initial pressure (bar)', '277 bar'),
        ('Temperature (T)', 'Reservoir temperature (°C)', '98 °C'),
        ('API Gravity', 'Stock tank oil API gravity', '39 °API'),
        ('GOR (Rs at Pi)', 'Solution gas-oil ratio at initial pressure', '160 m³/m³'),
        ('Water Salinity', 'Formation water salinity (ppm)', '50,000 ppm'),
    ],
    col_widths=[Cm(4.5), Cm(8.0), Cm(5.0)],
)
tip('Click Load Norne Defaults on the PVT step to populate all values from the '
    'Norne North Sea benchmark field.')

h2('5.5  Step 4 — Schedule')
body('Add production and injection control periods:')
bullet('Click + Add Period and enter the start date, end date, and control mode '
       '(RATE or BHP) for each period')
bullet('Periods are stored chronologically; each controls the target constraints '
       'for the PINO simulation')

h2('5.6  Step 5 — Review & Save')
body('Review the project summary, then:')
numbered('Enter a project name and choose a save location.')
numbered('Click Save Project. The file is written as ProjectName.pfproj (JSON format).')
numbered('Optionally enable Encrypt on Save and enter a password — see Section 13.')

body('After saving, PhysicsFlow registers the project in its SQLite database and returns '
     'you to the Dashboard.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. TRAINING THE PINO SURROGATE
# ══════════════════════════════════════════════════════════════════════════════

h1('6  Training the PINO Surrogate')

body('The Physics-Informed Neural Operator (PINO) surrogate is a 3D Fourier Neural Operator '
     '(FNO3d) trained to predict pressure and water saturation across the entire reservoir '
     'for a given set of permeability and porosity fields.')

h2('6.1  Training Data Requirements')
body('PINO training requires a dataset of input/output simulation pairs:')
bullet('Inputs: K (permeability), φ (porosity), P₀ (initial pressure), Sw₀ (initial water saturation), '
       'plus x and z coordinate channels — shape [N, 6, Nx, Ny, Nz]')
bullet('Targets: P(t) and Sw(t) for T timesteps — shape [N, T, 2, Nx, Ny, Nz]')
bullet('Minimum recommended: N = 200 samples; N = 500+ for production quality')

body('Training data can come from:')
bullet('OPM FLOW simulation runs (highest quality; 100–500 runs recommended)')
bullet('Synthetic ensemble generation via the physicsflow-pretrain CLI (see Section 14)')

h2('6.2  Starting Training from the UI')
numbered('Open your project and navigate to Training in the sidebar.')
numbered('Click Configure Training to set:')
bullet('Epochs — recommended 200 for first training, 100 for fine-tuning', level=2)
bullet('Batch size — 4 recommended for 12+ GB VRAM; 2 for smaller GPUs', level=2)
bullet('Learning rate — default 1e-3 with StepLR decay every 50 epochs', level=2)
bullet('PINO loss weights: w_data (1.0), w_pde (0.1), w_ic (0.5), w_bc (0.2), w_well (0.5)', level=2)
numbered('Click Start Training. Training progress is displayed in real time:')
bullet('Loss curves (total, PDE, data) plotted with OxyPlot', level=2)
bullet('Current epoch / total epochs counter', level=2)
bullet('Best checkpoint indicator with loss value', level=2)
numbered('Training automatically saves the best checkpoint to the models/ directory as '
         'pino_norne_pretrained.pt. Periodic checkpoints are saved every 50 epochs.')

warning('Do not close the application during training. Use Pause/Resume to interrupt safely. '
        'Closing mid-epoch will lose the current epoch\'s gradient updates.')

h2('6.3  PINO Architecture')
body('The FNO3d model architecture used by PhysicsFlow:')
simple_table(
    ['Parameter', 'Default', 'Description'],
    [
        ('Fourier modes (x, y)', '8', 'Number of Fourier modes in X and Y directions'),
        ('Fourier modes (z)', '6', 'Number of Fourier modes in Z direction'),
        ('Width', '32', 'Channel width of spectral convolution layers'),
        ('Layers', '4', 'Number of FNO blocks'),
        ('Input channels', '6', 'K_log, φ, P₀, Sw₀, x_coord, z_coord'),
        ('Output channels', '2', 'Pressure P and water saturation Sw'),
        ('Timesteps (T)', '20', 'Predicted timestep count'),
    ],
    col_widths=[Cm(5.0), Cm(3.0), Cm(9.5)],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. HISTORY MATCHING (αREKI)
# ══════════════════════════════════════════════════════════════════════════════

h1('7  History Matching (αREKI)')

body('History matching calibrates the reservoir model to match observed production data. '
     'PhysicsFlow uses Adaptive Regularised Ensemble Kalman Inversion (αREKI), '
     'a JAX-accelerated iterative Bayesian method.')

h2('7.1  Workflow Overview')
body('The history matching workflow is:')
numbered('Load observed production data (WOPR, WWPR, WGPR, WBHP) per well.')
numbered('Configure the ensemble (200 members recommended).')
numbered('Start αREKI — each iteration runs the PINO forward model for all ensemble members '
         'and updates the permeability/porosity fields via the Kalman gain.')
numbered('Monitor convergence via the mismatch plot in the History Matching view.')
numbered('Inspect the P10/P50/P90 envelope and per-well heatmap to assess match quality.')
numbered('Accept the calibrated ensemble or continue iterating.')

h2('7.2  Loading Observed Data')
bullet('Click Import Observations and select an Eclipse summary (.SMSPEC/.UNSMRY) or CSV file.')
bullet('Map each well\'s observed rates to the corresponding model wells.')
bullet('Optionally assign data weights — higher weight = that well drives calibration more strongly.')

h2('7.3  αREKI Configuration')
simple_table(
    ['Parameter', 'Default', 'Description'],
    [
        ('Ensemble size', '200', 'Number of realisations. 200 gives stable P10/P50/P90.'),
        ('Max iterations', '20', 'Maximum αREKI iterations before stopping.'),
        ('Localisation radius', '0.4', 'Gaspari-Cohn localisation radius (fraction of field extent).'),
        ('Noise inflation σ', '0.05', 'Observation noise standard deviation added to ensemble.'),
        ('Convergence target', 's_cumulative ≥ 1.0', 'Morozov discrepancy principle stopping criterion.'),
    ],
    col_widths=[Cm(4.5), Cm(3.0), Cm(10.0)],
)

h2('7.4  Understanding the Results')
body('The History Matching view shows:')
bullet('Ensemble fan chart — P10/P50/P90 bands vs. observed data for each well')
bullet('Mismatch convergence plot — normalised RMSE vs. αREKI iteration number')
bullet('Per-well RMSE heatmap — quickly identify wells with poor match')
bullet('α trace — adaptive step size per iteration; should converge smoothly to 0')

tip('If the mismatch stops converging after 5–8 iterations, try increasing the '
    'ensemble size to 500 or reducing the localisation radius.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. PRODUCTION FORECAST
# ══════════════════════════════════════════════════════════════════════════════

h1('8  Production Forecast')

body('After history matching, the calibrated ensemble is propagated forward in time to '
     'generate a probabilistic production forecast with P10/P50/P90 uncertainty bands.')

h2('8.1  Running a Forecast')
numbered('Navigate to Forecast in the sidebar.')
numbered('Confirm the forecast period start and end dates.')
numbered('Set the production control mode (RATE or BHP) for the forecast period.')
numbered('Click Run Forecast. The PINO surrogate runs all ensemble members in seconds.')
numbered('P10/P50/P90 fan charts are displayed for each well and for field totals.')

h2('8.2  EUR and Recovery Factor')
body('The Forecast view automatically computes:')
bullet('EUR (Estimated Ultimate Recovery) — P10/P50/P90 in MMSTB oil and BSCF gas')
bullet('Recovery Factor — EUR / STOIIP (requires STOIIP entered during project setup)')
bullet('Peak rate and peak rate date for oil, water, and gas')
bullet('Cumulative production tables')

h2('8.3  Exporting Results')
body('Two export formats are available from the Forecast view:')

h3('Export to PDF (QuestPDF)')
numbered('Click Export PDF → HM Summary Report to generate a report covering convergence '
         'history, per-well RMSE, and best ensemble member statistics.')
numbered('Click Export PDF → EUR Report to generate a report with P10/P50/P90 EUR table, '
         'per-well EUR breakdown, and regulatory disclaimer.')
numbered('Choose the save location when prompted. The PDF opens automatically after generation.')

h3('Export to Excel (ClosedXML)')
numbered('Click Export Excel to generate a multi-sheet workbook containing:')
bullet('Well Data — observed and simulated rates for every well at every timestep', level=2)
bullet('Ensemble Statistics — P10/P50/P90 statistics per well', level=2)
bullet('Training History — loss values per epoch for the active model', level=2)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 9. 3D RESERVOIR VIEWER
# ══════════════════════════════════════════════════════════════════════════════

h1('9  3D Reservoir Viewer')

body('The 3D Reservoir Viewer renders the full reservoir model as coloured voxels using '
     'HelixToolkit.Wpf. Navigate to Visualisation → 3D Viewer in the sidebar.')

h2('9.1  Controls')
simple_table(
    ['Control', 'Action'],
    [
        ('Left mouse drag', 'Rotate the model'),
        ('Right mouse drag', 'Pan the view'),
        ('Scroll wheel', 'Zoom in / out'),
        ('Middle mouse drag', 'Alternative pan'),
        ('Double-click', 'Reset view to fit model'),
    ],
    col_widths=[Cm(5.0), Cm(12.5)],
)

h2('9.2  Display Properties')
body('Use the toolbar controls to change what is rendered:')
bullet('Property selector — choose from Permeability (K_log), Porosity (φ), '
       'Pressure (P), Water Saturation (Sw), Gas Saturation (Sg)')
bullet('Timestep slider — scrub through simulation timesteps to animate the simulation')
bullet('Colormap selector — Jet (default), Viridis, or Grayscale')
bullet('Well markers toggle — show/hide yellow tube geometries at well locations')
bullet('Opacity slider — adjust cell opacity to see internal structure')

h2('9.3  Animation')
numbered('Select the property to animate (e.g., Pressure or Sw).')
numbered('Click Play to step through all timesteps automatically.')
numbered('Adjust playback speed with the FPS slider.')
numbered('Click Pause to stop at any timestep.')
numbered('Click Export VTK to save the current frame as a .vtu file for ResInsight or ParaView.')

note('The 3D viewer downsamples large grids for interactive display — every 4th cell in X '
     'and Y for the Norne 46×112×22 grid. Full-resolution VTK export is always available.')

h2('9.4  Well Visualisation')
body('Wells are displayed as yellow tube geometries:')
bullet('Producers are shown with a circle marker at the wellhead')
bullet('Injectors are shown with a diamond marker')
bullet('Hover over a well tube to see the well name and current WOPR/WWPR/WGPR in the tooltip')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 10. 2D CROSS-SECTION VIEWER
# ══════════════════════════════════════════════════════════════════════════════

h1('10  2D Cross-Section Viewer')

body('The 2D Cross-Section Viewer renders grid slices as pixel-accurate bitmaps using '
     'WriteableBitmap. Navigate to Visualisation → Cross-Section in the sidebar.')

h2('10.1  Slice Planes')
body('Three tabs provide views along each grid axis:')
bullet('I-Plane (YZ slice) — a plane of constant I index, showing Y × Z extent')
bullet('J-Plane (XZ slice) — a plane of constant J index, showing X × Z extent')
bullet('K-Plane (XY slice) — a layer slice at constant K depth, showing X × Y extent')

body('Use the slider on each tab to move the slice plane through the reservoir.')

h2('10.2  Display Properties')
simple_table(
    ['Control', 'Description'],
    [
        ('Property selector', 'K_log, Porosity, Pressure, Water Saturation, Gas Saturation'),
        ('Colormap', 'Jet, Viridis, Seismic (diverging), Greys'),
        ('Timestep', 'Select the simulation timestep (0 = static properties)'),
        ('Well overlay', 'Toggle the display of well perforations as coloured dots on the slice'),
        ('Colorbar', 'Min/max values displayed at the bottom of each slice image'),
    ],
    col_widths=[Cm(4.0), Cm(13.5)],
)

h2('10.3  Well Overlay')
body('When Well Overlay is enabled, well perforations that pass through the current slice plane '
     'are shown as coloured dots:')
bullet('Green dots — active producers')
bullet('Blue dots — active water injectors')
bullet('Red dots — active gas injectors')
bullet('Click a dot to open the well performance summary popup')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 11. REPORTS & EXCEL EXPORT
# ══════════════════════════════════════════════════════════════════════════════

h1('11  Reports & Excel Export')

h2('11.1  History Matching Summary Report (PDF)')
body('The HM Summary report provides a professional one-page-per-section document covering:')
bullet('Project identification: name, grid dimensions, ensemble size, convergence status')
bullet('Convergence history table: mismatch, α value, and P10/P50/P90 snapshot per iteration')
bullet('Per-well RMSE table: observed vs. simulated rates for all wells')
bullet('Best ensemble member statistics: lowest mismatch realisation summary')
numbered('From the Forecast view, click Export PDF → HM Summary Report.')
numbered('Select the output path when prompted.')
numbered('The PDF is generated using QuestPDF Community licence and opens automatically.')

h2('11.2  EUR Report (PDF)')
body('The EUR report is designed for reservoir management review:')
bullet('EUR P10/P50/P90 table: oil (MMSTB), water (MMSTB), gas (BSCF)')
bullet('Per-well EUR breakdown: individual well contributions')
bullet('Recovery factor: EUR / STOIIP with confidence bounds')
bullet('Regulatory disclaimer: standard statement for preliminary screening purposes')

warning('EUR reports are for internal screening only and are not certified for '
        'SEC or SPE-PRMS reserve bookings. Always consult a qualified reserves estimator '
        'before filing official reserve statements.')

h2('11.3  Excel Export')
body('The Excel workbook created by Export Excel contains three sheets:')

h3('Sheet 1: Well Data')
body('One row per well per timestep, with columns:')
code_block('WellName | Timestep | ObsOilRate | SimOilRate | ObsWaterRate | SimWaterRate |\n'
           'ObsGasRate | SimGasRate | ObsBHP | SimBHP | OilP10 | OilP90 | WaterP10 | WaterP90')

h3('Sheet 2: Ensemble Statistics')
body('P10/P50/P90 statistics per well for oil, water, and gas rates at each timestep.')

h3('Sheet 3: Training History')
body('Loss values per training epoch:')
code_block('Epoch | LossTotal | LossPDE | LossData | LossWell | LossIC | LossBC | LearningRate')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 12. AI RESERVOIR ASSISTANT
# ══════════════════════════════════════════════════════════════════════════════

h1('12  AI Reservoir Assistant')

body('The AI Reservoir Assistant is an embedded local LLM (powered by Ollama) grounded by '
     'three complementary intelligence layers — all running entirely on your machine:')
bullet('Layer 1 — Live tool calls: 8 tools that read your real-time simulation data')
bullet('Layer 2 — Reservoir Knowledge Graph: structured facts about Norne topology, wells, faults, and parameters')
bullet('Layer 3 — Hybrid RAG: indexed project documents, well reports, and technical papers')
body('No data is sent to the internet. All inference, embeddings, and vector search run locally.')

h2('12.1  Starting the Assistant')
body('The AI panel is always visible on the right side of the screen. If Ollama is not '
     'running, a prompt at the top of the panel will ask you to start it.')
numbered('Ensure Ollama is running: open a terminal and type ollama list.')
numbered('If needed, click the "Start Ollama" button in the AI panel.')
numbered('Type your question in the chat input at the bottom of the panel.')
numbered('Press Enter or click Send. The assistant will respond with streaming tokens.')

h2('12.2  Available Tools')
body('The AI assistant has access to 10 tools across the three intelligence layers:')
simple_table(
    ['Tool', 'Layer', 'What it returns'],
    [
        ('get_simulation_status', 'Live', 'Current training epoch, loss, model type, and training duration'),
        ('get_well_performance', 'Live', 'WOPR/WWPR/WGPR rates for a named well at the latest timestep'),
        ('get_hm_summary', 'Live', 'Current αREKI iteration, mismatch value, α, and convergence status'),
        ('get_ensemble_stats', 'Live', 'P10/P50/P90 statistics for a well or field total'),
        ('get_field_property', 'Live', 'Mean/min/max of a reservoir property (K, φ, P, Sw) for any layer'),
        ('explain_parameter', 'Live', 'Plain-English explanation of any parameter in the project file'),
        ('get_training_history', 'Live', 'Full epoch-by-epoch loss history for the current model'),
        ('get_well_list', 'Live', 'List all wells with type (producer/injector) and perforation depths'),
        ('query_reservoir_graph', 'KG', '20-pattern NL query over reservoir topology — wells, layers, faults, segments'),
        ('search_project_knowledge', 'RAG', 'Hybrid semantic + keyword search over all indexed project documents'),
    ],
    col_widths=[Cm(4.5), Cm(1.5), Cm(11.5)],
)

h2('12.3  Example Questions — Live Data')
body('The AI assistant works best with specific, data-grounded questions:')
bullet('"What is the current WOPR for well E-1H?"')
bullet('"How many αREKI iterations have been completed and what is the current mismatch?"')
bullet('"What is the P50 cumulative oil production forecast for the field?"')
bullet('"Explain what the w_pde loss weight controls in the PINO training."')
bullet('"Which wells have the highest RMSE in the history match?"')
bullet('"What is the average permeability in layer 10 of the reservoir?"')

h2('12.4  Example Questions — Knowledge Graph (Structural)')
body('The Reservoir Knowledge Graph is pre-populated with Norne structural facts. Ask:')
bullet('"Which layers does well B-2H perforate?"')
bullet('"Which injectors support producers in segment C?"')
bullet('"What segments are connected to segment B?"')
bullet('"Which parameters influence water cut (WWCT)?"')
bullet('"What faults bound segment D?"')
bullet('"List all producers in segment E."')
body('These queries are answered directly from the graph — zero latency, fully deterministic.')

h2('12.5  Example Questions — Knowledge Base (RAG)')
body('After indexing project documents (see Chapter 13), the assistant can answer '
     'document-grounded questions:')
bullet('"What does the Norne field study say about the northern fault block?"')
bullet('"What is the OWC depth for segment B according to the well completion report?"')
bullet('"Summarise the key uncertainties identified in the subsurface evaluation report."')
tip('Index your project documents in the Knowledge Base Management panel (Chapter 13) '
    'before asking document-based questions.')

h2('12.6  Changing the Model')
body('The default model is phi3:mini (3.8B parameters — fast, good for question answering). '
     'To use a different model:')
numbered('Pull the model in Ollama: ollama pull llama3:8b')
numbered('In PhysicsFlow, go to Settings → AI Assistant → Model.')
numbered('Select the model from the dropdown and click Apply.')
note('A larger model like llama3:8b gives better analytical responses but requires 8+ GB VRAM. '
     'phi3:mini runs well on CPU with 8 GB RAM and is recommended for most users.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 13. KNOWLEDGE BASE MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

h1('13  Knowledge Base Management (RAG + Knowledge Graph)')

body('PhysicsFlow v1.3.0 includes a Hybrid RAG (Retrieval-Augmented Generation) system '
     'and a pre-populated Reservoir Knowledge Graph. This chapter explains how to '
     'index your project documents and query the knowledge base.')

h2('13.1  How the Intelligence Layer Works')
body('When you ask the AI a question, PhysicsFlow runs three grounding queries in parallel '
     'before calling the LLM:')
numbered('Knowledge Graph lookup — regex-matched NL query over the networkx graph')
numbered('Hybrid RAG retrieval — ChromaDB dense search + BM25 sparse search, fused with RRF, '
         'then reranked by a cross-encoder (MiniLM)')
numbered('Live tool calls — real-time simulation state from the Python engine')
body('The KG answer is injected first (highest precision), then the RAG context, then the '
     'live data. The LLM synthesises a grounded response from all three sources.')

h2('13.2  Indexing Documents')
body('Supported file types: PDF, Word (.docx), plain text (.txt), CSV, LAS 2.0, Eclipse .DATA. '
     'To index documents:')
numbered('In PhysicsFlow, navigate to Settings → Knowledge Base.')
numbered('Click Add Files or Add Directory to select documents.')
numbered('PhysicsFlow will chunk, embed, and index all documents. A progress bar shows status.')
numbered('Once indexed, the chunk count is shown next to each source file.')

body('Alternatively, index documents from the Python API:')
code_block('from physicsflow.rag import RAGPipeline\n\n'
           'rag = RAGPipeline.instance()\n'
           'rag.index_file("reports/Norne_field_study.pdf")     # PDF\n'
           'rag.index_file("data/B-2H.las")                     # LAS 2.0 well log\n'
           'rag.index_directory("docs/")                        # entire folder\n'
           'print(f"Indexed: {rag.count()} chunks")')

tip('Index the field development plan, well completion reports, and any reference papers '
    'relevant to your study. The RAG system retrieves the most relevant passages automatically.')

h2('13.3  Managing Indexed Documents')
body('In Settings → Knowledge Base, you can:')
bullet('View all indexed sources with chunk count and last-indexed timestamp')
bullet('Remove a source — deletes all chunks from that file')
bullet('Re-index a source — use this after the source document is updated')
bullet('Clear all — wipes the entire vector store (irreversible)')
bullet('View chunk count and storage path (%APPDATA%\\PhysicsFlow\\rag\\chroma\\)')

h2('13.4  The Reservoir Knowledge Graph')
body('The Reservoir Knowledge Graph (KG) is automatically built on first startup from:')
simple_table(
    ['Source', 'What it adds'],
    [
        ('Norne structural base', '22 layers (K1-K22), 17 producers, 5 injectors, 53 faults, 5 segments (A-E), segment connectivity, injector-producer support pairs, uncertain parameters'),
        ('.pfproj project file', 'Additional wells from project file, perforation depths, HM results'),
        ('SQLite database', 'Simulation runs, HM iterations, convergence status, per-run mismatch'),
        ('Live context provider', 'Per-well RMSE from the last HM run — updated after each iteration'),
    ],
    col_widths=[Cm(4.5), Cm(13.0)],
)
body('The KG graph is persisted to disk (%APPDATA%\\PhysicsFlow\\kg\\reservoir_graph.json) '
     'and reloaded on startup. Call PhysicsFlow.rebuild_kg() to force a full rebuild.')

h2('13.5  Supported Natural Language KG Query Patterns')
body('The KG query engine recognises 20 natural language patterns:')
simple_table(
    ['Example Query', 'What it returns'],
    [
        ('"which layers does B-2H perforate"', 'Perforation layers of the well (sorted K1..K22)'),
        ('"wells in layer K10"', 'All wells with completions in that layer'),
        ('"wells in segment C"', 'All wells (producer + injector) in segment C'),
        ('"which segment is E-1H in"', 'Segment assignment of the well'),
        ('"which injectors support D-2H"', 'Injector wells providing pressure support'),
        ('"producers supported by F-4H"', 'Producer wells supported by the injector'),
        ('"faults in segment B"', 'Fault names bounding a segment'),
        ('"which segments does F-NW bound"', 'Segments on either side of the fault'),
        ('"connected segments to A"', 'Graph-reachable segment neighbours'),
        ('"parameters influencing WOPR"', 'Uncertain parameters that affect the quantity'),
        ('"which runs converged"', 'Simulation runs with convergence flag'),
        ('"list all producers"', 'Full producer well list'),
        ('"list all injectors"', 'Full injector well list'),
        ('"what type is well C-4AH"', 'Producer / injector / unknown'),
        ('"worst matching wells"', 'Wells sorted descending by last_rmse'),
    ],
    col_widths=[Cm(7.0), Cm(10.5)],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 14. PROJECT FILE ENCRYPTION
# ══════════════════════════════════════════════════════════════════════════════

h1('14  Project File Encryption')

body('PhysicsFlow project files (.pfproj) contain sensitive reservoir data including '
     'permeability fields, well data, and production history. AES-256-GCM encryption '
     'is available to protect this data at rest.')

h2('14.1  Encrypting a Project File (UI)')
numbered('Open your project in PhysicsFlow.')
numbered('Click File → Save As Encrypted, or enable Encrypt on Save in the Project Setup wizard.')
numbered('Enter a strong password when prompted. Enter it again to confirm.')
numbered('The project is saved as ProjectName.pfproj.enc.')
numbered('The original .pfproj file is securely deleted if you selected Remove Original.')

note('The encrypted file uses a .pfproj.enc extension. PhysicsFlow automatically detects '
     'this extension and prompts for the password when opening.')

h2('14.2  Encryption Technical Details')
body('PhysicsFlow uses industry-standard authenticated encryption:')
simple_table(
    ['Property', 'Value'],
    [
        ('Algorithm', 'AES-256-GCM (authenticated encryption)'),
        ('Key derivation', 'PBKDF2-HMAC-SHA256'),
        ('Iterations', '600,000 (NIST recommended minimum for 2026)'),
        ('Salt', '32 bytes, randomly generated per encryption'),
        ('Nonce', '12 bytes, randomly generated per encryption'),
        ('Authentication tag', '16 bytes (GCM tag — detects tampering)'),
        ('File format', 'PFEC binary: magic(4) + version(1) + iters(4) + salt(32) + nonce(12) + tag(16) + ciphertext'),
    ],
    col_widths=[Cm(4.5), Cm(13.0)],
)

h2('14.3  Encrypting via the CLI')
body('The physicsflow-encrypt and physicsflow-decrypt commands are available in the terminal:')

code_block('# Encrypt a project file (prompts for password)\nphysicsflow-encrypt study.pfproj\n\n'
           '# Encrypt with explicit output path and securely delete original\nphysicsflow-encrypt study.pfproj -o study_enc.pfproj.enc --remove-original\n\n'
           '# Decrypt a project file\nphysicsflow-decrypt study.pfproj.enc\n\n'
           '# Decrypt to a specific output path\nphysicsflow-decrypt study.pfproj.enc -o /path/to/output.pfproj')

warning('There is no password recovery mechanism. If you forget the encryption password, '
        'the project file cannot be decrypted. Keep a secure copy of the password.')

h2('14.4  Secure Delete')
body('When Remove Original is selected, PhysicsFlow performs a best-effort secure delete:')
bullet('Overwrites the file with zeros (single pass)')
bullet('Calls fsync to flush to disk')
bullet('Unlinks the file from the filesystem')

note('Single-pass overwrite is sufficient for SSD and HDD protection against casual '
     'data recovery. For high-security environments, use full-disk encryption (BitLocker) '
     'in addition to project file encryption.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 14. PINO PRE-TRAINING CLI
# ══════════════════════════════════════════════════════════════════════════════

h1('15  PINO Pre-Training CLI')

body('The physicsflow-pretrain command generates a synthetic Norne ensemble and trains '
     'the FNO3d PINO surrogate. This is useful for:')
bullet('Bootstrapping a new project before OPM FLOW training data is available')
bullet('Testing the training pipeline and GPU configuration')
bullet('Generating a starting checkpoint for transfer learning to a new field')

h2('15.1  Basic Usage')
code_block('physicsflow-pretrain --epochs 200 --ensemble 500 --device cuda')

h2('15.2  All Options')
simple_table(
    ['Option', 'Default', 'Description'],
    [
        ('--deck', 'None', 'Path to NORNE_ATW2013.DATA Eclipse deck (optional; uses synthetic K/φ if omitted)'),
        ('--output-dir', 'models/', 'Directory to save checkpoints'),
        ('--epochs', '200', 'Number of training epochs'),
        ('--ensemble', '500', 'Number of ensemble members (synthetic K/φ perturbations)'),
        ('--batch-size', '4', 'Batch size per gradient step'),
        ('--lr', '1e-3', 'Initial learning rate (StepLR decays by 0.5 every 50 epochs)'),
        ('--device', 'cuda', "'cuda' or 'cpu'"),
        ('--seed', '42', 'Random seed for reproducibility'),
        ('--width', '32', 'FNO channel width'),
        ('--modes', '8', 'Fourier modes in X and Y'),
        ('--modes-z', '6', 'Fourier modes in Z'),
        ('--w-pde', '0.1', 'PDE loss weight'),
        ('--w-data', '1.0', 'Data loss weight'),
        ('--log-every', '10', 'Print log every N epochs'),
        ('--save-every', '50', 'Save checkpoint every N epochs'),
    ],
    col_widths=[Cm(3.5), Cm(2.5), Cm(11.5)],
)

h2('15.3  Output Files')
body('The command saves the following files in the output directory:')
bullet('pino_norne_pretrained.pt — best checkpoint (lowest total loss)')
bullet('pino_norne_epoch0050.pt, pino_norne_epoch0100.pt, … — periodic checkpoints')
body('Each checkpoint contains the model state dict, epoch number, loss value, and config dictionary.')
body('The training run is also registered in the PhysicsFlow SQLite database (training_epochs table).')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 15. SETTINGS & CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

h1('16  Settings & Configuration')

h2('16.1  Application Settings (UI)')
body('Access Settings from the gear icon in the bottom-left of the sidebar:')

simple_table(
    ['Setting', 'Description', 'Default'],
    [
        ('Python Engine Port', 'gRPC port for Python engine communication', '50051'),
        ('Engine Startup Timeout', 'Seconds to wait for engine.ready before error', '60 s'),
        ('GPU Device', 'CUDA device index (0 = first GPU)', '0'),
        ('Ollama Endpoint', 'URL of local Ollama server', 'http://localhost:11434'),
        ('AI Model', 'Ollama model name for AI assistant', 'phi3:mini'),
        ('Database Path', 'SQLite database file location', '%APPDATA%\\PhysicsFlow\\physicsflow.db'),
        ('Log Level', 'Serilog log level (Verbose / Debug / Info / Warning)', 'Information'),
        ('Theme', 'MahApps.Metro dark/light theme', 'Dark'),
        ('Default Output Dir', 'Default directory for exports and reports', 'Documents\\PhysicsFlow'),
    ],
    col_widths=[Cm(4.5), Cm(9.0), Cm(4.0)],
)

h2('16.2  Environment Variables (Python Engine)')
body('The Python engine reads configuration from PHYSICSFLOW_* environment variables:')
code_block('PHYSICSFLOW_PORT=50051           # gRPC server port\n'
           'PHYSICSFLOW_DB_PATH=<path>       # Override SQLite path\n'
           'PHYSICSFLOW_LOG_LEVEL=INFO       # Loguru level\n'
           'PHYSICSFLOW_DEVICE=cuda          # torch.device string\n'
           'PHYSICSFLOW_OLLAMA_URL=http://localhost:11434')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 16. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

h1('17  Troubleshooting')

simple_table(
    ['Problem', 'Likely Cause', 'Solution'],
    [
        ('Engine fails to start',
         'Python not found or gRPC port in use',
         'Check port 50051 is free (netstat -an | findstr 50051). '
         'Reinstall the Python components via Settings → Repair Engine.'),
        ('"Engine not connected" in status bar',
         'Engine.ready signal file not created within timeout',
         'Increase the engine startup timeout in Settings. Check engine logs in '
         '%APPDATA%\\PhysicsFlow\\logs\\engine.log for Python errors.'),
        ('CUDA out of memory during training',
         'GPU VRAM insufficient for batch size',
         'Reduce batch size to 2 or 1. Reduce FNO width from 32 to 16.'),
        ('Slow training (CPU mode)',
         'No NVIDIA GPU, or CUDA not found',
         'Verify CUDA availability: python -c "import torch; print(torch.cuda.is_available())"'),
        ('Cannot open .pfproj.enc',
         'Wrong password or file corrupted',
         'Verify you are using the correct password. '
         'The error "file has been tampered with" indicates data corruption.'),
        ('Ollama not responding',
         'Ollama service not running',
         'Open a terminal and run: ollama serve. '
         'Then click Reconnect in the AI panel.'),
        ('3D viewer blank / no voxels',
         'No simulation data loaded',
         'Run a forward simulation or load a project with training data first.'),
        ('Excel export empty sheets',
         'No history matching results',
         'Complete at least one αREKI iteration before exporting well data.'),
        ('PDF report "Could not create file"',
         'Output directory permissions',
         'Choose a writable output directory such as Documents\\PhysicsFlow\\Reports.'),
    ],
    col_widths=[Cm(4.0), Cm(5.5), Cm(8.0)],
)

h2('17.1  Log Files')
body('PhysicsFlow writes detailed logs to:')
bullet('%APPDATA%\\PhysicsFlow\\logs\\physicsflow.log — .NET desktop application log (Serilog)')
bullet('%APPDATA%\\PhysicsFlow\\logs\\engine.log — Python gRPC engine log (Loguru)')
body('Set Log Level to Verbose in Settings to capture maximum detail for bug reports.')

h2('17.2  Reporting Issues')
body('Report bugs or feature requests at:')
code_block('https://github.com/Danny024/PhysicsFlow/issues')
body('Please include:')
bullet('PhysicsFlow version (visible in Help → About)')
bullet('Windows version and GPU model')
bullet('The contents of physicsflow.log and engine.log (redact any sensitive field data)')
bullet('Steps to reproduce the problem')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 17. KEYBOARD SHORTCUTS
# ══════════════════════════════════════════════════════════════════════════════

h1('18  Keyboard Shortcuts')

simple_table(
    ['Shortcut', 'Action'],
    [
        ('Ctrl+N', 'New Project'),
        ('Ctrl+O', 'Open Project'),
        ('Ctrl+S', 'Save Project'),
        ('Ctrl+Shift+S', 'Save Project As'),
        ('Ctrl+E', 'Export PDF'),
        ('Ctrl+Shift+E', 'Export Excel'),
        ('F5', 'Start / Resume Training'),
        ('F6', 'Pause Training'),
        ('F7', 'Start History Matching'),
        ('F8', 'Run Forecast'),
        ('Ctrl+3', 'Navigate to 3D Viewer'),
        ('Ctrl+2', 'Navigate to 2D Cross-Section'),
        ('Ctrl+D', 'Navigate to Dashboard'),
        ('Ctrl+T', 'Navigate to Training'),
        ('Ctrl+H', 'Navigate to History Matching'),
        ('Ctrl+F', 'Navigate to Forecast'),
        ('Ctrl+/', 'Focus AI Assistant chat input'),
        ('Esc', 'Cancel current operation / close dialog'),
        ('F1', 'Open User Manual (this document)'),
    ],
    col_widths=[Cm(4.0), Cm(13.5)],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 18. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════

h1('19  Glossary')

simple_table(
    ['Term', 'Definition'],
    [
        ('AES-256-GCM', 'Advanced Encryption Standard, 256-bit key, Galois/Counter Mode — authenticated encryption algorithm'),
        ('αREKI', 'Adaptive Regularised Ensemble Kalman Inversion — iterative Bayesian history matching method'),
        ('BHP', 'Bottom-hole pressure — pressure measured at the perforated interval of a well'),
        ('CCR', 'Cluster-Classify-Regress — XGBoost mixture-of-experts well surrogate model'),
        ('DDIM', 'Denoising Diffusion Implicit Model — fast sampling variant of diffusion generative model'),
        ('EUR', 'Estimated Ultimate Recovery — total hydrocarbon volume expected to be produced over the field life'),
        ('FNO', 'Fourier Neural Operator — neural network architecture for learning solution operators of PDEs'),
        ('GCM', 'Galois/Counter Mode — block cipher mode providing both encryption and authentication'),
        ('GOR', 'Gas-Oil Ratio — volume of gas produced per volume of oil at surface conditions'),
        ('K', 'Permeability — measure of a rock\'s ability to transmit fluid (unit: millidarcy, mD)'),
        ('PBKDF2', 'Password-Based Key Derivation Function 2 — standard for deriving a cryptographic key from a password'),
        ('PFEC', 'PhysicsFlow Encrypted Container — binary file format for encrypted .pfproj files'),
        ('pfproj', 'PhysicsFlow project file — JSON document storing all study configuration'),
        ('PINO', 'Physics-Informed Neural Operator — neural operator trained with PDE residual loss'),
        ('P10/P50/P90', 'Percentile estimates for an uncertain quantity — P50 is the median; P10/P90 bound the 80% confidence interval'),
        ('PVT', 'Pressure-Volume-Temperature — fluid property correlations (μ, Rs, Bo, Bg, Bw)'),
        ('RMSE', 'Root Mean Square Error — measure of mismatch between observed and simulated values'),
        ('STOIIP', 'Stock Tank Oil Initially In Place — total volume of oil in the reservoir at initial conditions'),
        ('UQ', 'Uncertainty Quantification — characterisation of the range and probability of possible outcomes'),
        ('VCAE', 'Variational Convolutional AutoEncoder — encodes permeability fields to a compact latent space'),
        ('VTK', 'Visualization Toolkit — file format supported by ResInsight, ParaView, and other 3D tools'),
        ('WAL', 'Write-Ahead Logging — SQLite journal mode enabling concurrent reads while writing'),
        ('WOPR', 'Well Oil Production Rate — surface oil rate for a producer well (m³/day or STB/day)'),
        ('WWPR', 'Well Water Production Rate — surface water rate for a producer well'),
        ('WGPR', 'Well Gas Production Rate — surface gas rate for a producer well'),
        ('RAG', 'Retrieval-Augmented Generation — technique that grounds LLM responses with retrieved document passages'),
        ('BM25', 'Best Match 25 — probabilistic sparse keyword ranking algorithm for document retrieval'),
        ('Cross-Encoder', 'A reranking model that jointly scores a query-document pair for precision; used after BM25/dense retrieval'),
        ('HyDE', 'Hypothetical Document Embedding — query expansion technique that generates a synthetic answer then uses its embedding for retrieval'),
        ('Knowledge Graph (KG)', 'Structured graph of reservoir entities (wells, layers, faults, segments) with typed edges representing geological relationships'),
        ('RRF', 'Reciprocal Rank Fusion — score fusion algorithm that combines rankings from dense and sparse retrieval without re-scoring'),
        ('ChromaDB', 'Open-source vector database for storing and querying dense embeddings; used as the PhysicsFlow RAG vector store'),
        ('BGE', 'BAAI/bge-small-en-v1.5 — 33M-parameter sentence embedding model from BAAI, used for RAG document embeddings'),
    ],
    col_widths=[Cm(3.5), Cm(14.0)],
)

doc.add_paragraph()
body('— End of PhysicsFlow User Manual v1.3.0 —', italic=True, color=MID_GREY)


# ── Save ───────────────────────────────────────────────────────────────────────
output_path = 'PhysicsFlow_UserManual_v130.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
