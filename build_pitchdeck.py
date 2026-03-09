"""
PhysicsFlow Pitch Deck — Word Document Generator
Creates a professionally structured, properly aligned pitch deck in .docx format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x0A, 0x16, 0x28)
BLUE       = RGBColor(0x15, 0x65, 0xC0)
CYAN       = RGBColor(0x00, 0xBC, 0xD4)
GREEN      = RGBColor(0x00, 0xC8, 0x53)
AMBER      = RGBColor(0xFF, 0x8F, 0x00)
RED        = RGBColor(0xE5, 0x39, 0x35)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_BG    = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_GREY = RGBColor(0xF0, 0xF4, 0xF8)
MID_GREY   = RGBColor(0x78, 0x90, 0x9C)
DARK_GREY  = RGBColor(0x37, 0x47, 0x4F)
SLIDE_BG   = RGBColor(0x0D, 0x1F, 0x3C)


# ── Helper: set paragraph background shading ──────────────────────────────────
def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 6)))
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_page_background(doc, hex_color: str):
    background = OxmlElement('w:background')
    background.set(qn('w:color'), hex_color)
    doc.element.insert(0, background)
    settings = doc.settings.element
    ds = OxmlElement('w:displayBackgroundShape')
    settings.append(ds)


def cell_vertical_align(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page size: A4 landscape
section = doc.sections[0]
section.page_width  = Cm(33.87)   # A4 landscape width
section.page_height = Cm(19.05)   # A4 landscape height
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)
section.top_margin    = Cm(1.4)
section.bottom_margin = Cm(1.4)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_GREY

CONTENT_WIDTH = section.page_width - section.left_margin - section.right_margin


# ── Utility functions ─────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=NAVY, size=28, bold=True, space_before=0, space_after=6, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    return p


def add_body(doc, text, color=DARK_GREY, size=10, bold=False, italic=False,
             space_before=0, space_after=4, center=False, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_label(doc, text, color=CYAN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.font.name   = 'Calibri'
    run.font.size   = Pt(8)
    run.font.bold   = True
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, color=DARK_GREY, size=10, bullet='•', indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run(f'{bullet}  {text}')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p


def add_divider(doc, color='0A1628'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Oxml
    br = _Oxml('w:br')
    br.set(_qn('w:type'), 'page')
    return br


def add_page_break(doc):
    doc.add_page_break()


def make_table(doc, rows, cols, col_widths=None, style='Table Grid'):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = style
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                cell.width = col_widths[j]
    return tbl


def cell_para(cell, text, bold=False, size=10, color=DARK_GREY, center=False, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_slide_header(doc, slide_num, total, section_label, label_color=CYAN):
    """Adds slide number + section label header row."""
    tbl = make_table(doc, 1, 2, col_widths=[CONTENT_WIDTH * 0.7, CONTENT_WIDTH * 0.3])
    shade_cell(tbl.cell(0, 0), '0A1628')
    shade_cell(tbl.cell(0, 1), '0A1628')
    cell_vertical_align(tbl.cell(0, 0), 'center')
    cell_vertical_align(tbl.cell(0, 1), 'center')
    lp = cell_para(tbl.cell(0, 0), section_label.upper(), bold=True, size=8, color=label_color)
    rp = tbl.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run(f'{slide_num} / {total}')
    run.font.size = Pt(8)
    run.font.color.rgb = MID_GREY
    run.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 1: COVER
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'PhysicsFlow', level=1, color=CYAN, size=52, bold=True,
            space_before=20, space_after=2, center=True)

add_body(doc, 'The AI-Native Reservoir Simulation Platform',
         color=WHITE, size=16, space_before=0, space_after=10, center=True)

add_divider(doc, color='1565C0')

add_body(doc,
    'Physics-Informed Neural Operator  ·  Ensemble History Matching  ·  Built-in AI Assistant',
    color=CYAN, size=11, italic=True, space_before=4, space_after=16, center=True)

# Key stats table
stats = make_table(doc, 2, 3,
    col_widths=[CONTENT_WIDTH//3, CONTENT_WIDTH//3, CONTENT_WIDTH//3])
stats.style = 'Table Grid'

data = [
    ('6,000×', 'FASTER THAN ECLIPSE 100'),
    ('~1 Hour', 'FULL HISTORY MATCH'),
    ('90%', 'COST REDUCTION VS. HPC'),
]
for j, (num, label) in enumerate(data):
    shade_cell(stats.cell(0, j), '0D2B5E')
    shade_cell(stats.cell(1, j), '0A1628')
    cell_para(stats.cell(0, j), num, bold=True, size=28, color=CYAN, center=True)
    cell_para(stats.cell(1, j), label, bold=True, size=8, color=MID_GREY, center=True)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

add_body(doc, 'Seed Round — $3.5M  |  contact@physicsflow.ai  |  physicsflow.ai',
         color=MID_GREY, size=9, space_before=8, space_after=0, center=True)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 2: THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 2, 16, '⚠  THE PROBLEM', RED)
add_heading(doc, 'Reservoir Engineering Is Broken', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'The tools have not changed fundamentally in 40 years — and the industry is paying for it.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=10)

# Three pain columns
pain = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2)])

pain_data = [
    ('🐌  Unacceptably Slow', '2–8 hours per simulation run. History matching requires thousands of runs. Engineers wait months for results that should take hours.', '6 months', 'Average HM project duration'),
    ('💸  Prohibitively Expensive', 'Eclipse licenses: $150,000–$400,000/year. HPC clusters: $500,000–$5M. Only supermajors can afford proper simulation workflows.', '$2M+', 'Cost per HM project'),
    ('🎲  No Real Uncertainty', 'Most operators run one deterministic model and call it "the" reservoir. No P10/P50/P90. Billions in capex allocated based on a single guess.', '85%', 'Of HM projects deliver one deterministic model only'),
]

for j, (title, desc, stat, stat_label) in enumerate(pain_data):
    shade_cell(pain.cell(0, j), 'FFF8F8')
    cell_vertical_align(pain.cell(0, j), 'top')
    c = pain.cell(0, j)
    # Title
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RED; r.font.name = 'Calibri'
    # Description
    p2 = c.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(6)
    p2.runs[0].font.size = Pt(9.5); p2.runs[0].font.name = 'Calibri'; p2.runs[0].font.color.rgb = DARK_GREY
    # Stat
    p3 = c.add_paragraph()
    r3 = p3.add_run(stat)
    r3.font.size = Pt(26); r3.font.bold = True; r3.font.color.rgb = RED; r3.font.name = 'Calibri'
    p4 = c.add_paragraph(stat_label)
    p4.runs[0].font.size = Pt(8); p4.runs[0].font.color.rgb = MID_GREY; p4.runs[0].font.name = 'Calibri'
    p4.paragraph_format.space_after = Pt(6)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Hidden cost callout
co = make_table(doc, 1, 1, col_widths=[CONTENT_WIDTH])
shade_cell(co.cell(0, 0), 'FFF3F3')
c = co.cell(0, 0)
p = c.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run('⚠  The Hidden Cost of Wrong Models:  ')
r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RED; r.font.name = 'Calibri'
r2 = p.add_run(
    'A 5% improvement in recovery factor on a 500 MMBOE field is worth $1.5 billion at $30/barrel. '
    'Poorly-matched reservoir models misallocate drilling capex, delay development decisions, '
    'and destroy value every single day.')
r2.font.size = Pt(10); r2.font.color.rgb = DARK_GREY; r2.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 3: THE SOLUTION
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 3, 16, '✅  THE SOLUTION', CYAN)
add_heading(doc, 'PhysicsFlow: The AI-Native Reservoir Platform', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Three breakthroughs combined into one industrial-grade desktop application.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=10)

sol = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2)])

sol_data = [
    ('🧠  PINO Surrogate', 'E8F4FD', BLUE,
     'A Physics-Informed Neural Operator trained on real OPM FLOW simulation data. Predicts pressure, '
     'water and gas saturation fields across the entire reservoir at every time step. Enforces Darcy '
     'PDE physics through the loss function — not just curve fitting.',
     '~7 seconds', 'per full simulation (vs. 3–8 hours Eclipse)'),
    ('🎯  αREKI History Matching', 'F0FFF4', GREEN,
     'Adaptive Regularised Ensemble Kalman Inversion automatically calibrates 200 reservoir realisations '
     'to match observed production data. JAX-accelerated ensemble operations deliver fully calibrated '
     'P10/P50/P90 uncertainty quantification in approximately one hour.',
     '200 members', 'full ensemble UQ in ~1 hour'),
    ('🤖  AI Reservoir Assistant', 'FFFBF0', AMBER,
     'An embedded local LLM (Ollama — runs on your machine, data never leaves) with live tool-calling '
     'access to simulation results. Engineers ask questions in plain English and get data-grounded '
     'answers instantly. No hallucinations: every answer is backed by a tool call.',
     'Local AI', 'data stays on-premise, fully private'),
]

for j, (title, bg, col, desc, stat, stat_label) in enumerate(sol_data):
    shade_cell(sol.cell(0, j), bg)
    cell_vertical_align(sol.cell(0, j), 'top')
    c = sol.cell(0, j)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = col; r.font.name = 'Calibri'
    p2 = c.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(8)
    p2.runs[0].font.size = Pt(9.5); p2.runs[0].font.color.rgb = DARK_GREY; p2.runs[0].font.name = 'Calibri'
    p3 = c.add_paragraph()
    r3 = p3.add_run(stat)
    r3.font.size = Pt(22); r3.font.bold = True; r3.font.color.rgb = col; r3.font.name = 'Calibri'
    p4 = c.add_paragraph(stat_label)
    p4.runs[0].font.size = Pt(8); p4.runs[0].font.color.rgb = MID_GREY; p4.runs[0].font.name = 'Calibri'
    p4.paragraph_format.space_after = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Workflow flow
add_body(doc, 'Workflow:', color=NAVY, size=10, bold=True, space_before=0, space_after=4)
flow = make_table(doc, 1, 9, col_widths=[
    Cm(3.3), Cm(0.6), Cm(3.3), Cm(0.6), Cm(3.3), Cm(0.6), Cm(3.3), Cm(0.6), Cm(3.3)])
flow_items = ['📋 Eclipse Deck\nImport', '→', '🔬 OPM FLOW\n100 training runs',
              '→', '🧠 PINO Training\n5hrs on GPU', '→',
              '⚡ αREKI HM\n~1 hour', '→', '📊 P10/P50/P90\nCalibrated UQ']
flow_colors = ['D6EAF8', 'FFFFFF', 'D6EAF8', 'FFFFFF', 'D6EAF8', 'FFFFFF', 'D5F5E3', 'FFFFFF', 'D5F5E3']
for j, (item, bg) in enumerate(zip(flow_items, flow_colors)):
    shade_cell(flow.cell(0, j), bg)
    cell_vertical_align(flow.cell(0, j), 'center')
    is_arrow = item == '→'
    cp = cell_para(flow.cell(0, j), item, bold=not is_arrow, size=9,
                   color=MID_GREY if is_arrow else DARK_GREY, center=True)
    flow.cell(0, j).paragraphs[0].paragraph_format.space_before = Pt(4)
    flow.cell(0, j).paragraphs[0].paragraph_format.space_after  = Pt(4)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 4: HOW IT WORKS
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 4, 16, '⚙  TECHNOLOGY OVERVIEW', CYAN)
add_heading(doc, 'How PhysicsFlow Works', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Hybrid PyTorch + JAX engine — best framework for each phase of the workflow.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=10)

tech = make_table(doc, 1, 2,
    col_widths=[CONTENT_WIDTH // 2 - Cm(0.3), CONTENT_WIDTH // 2 - Cm(0.3)])

tech_left = [
    ('Training Phase — PyTorch + PhysicsNeMo (NVIDIA)', BLUE, [
        ('📐  Fourier Neural Operator (FNO)',
         'Three separate models: pressure, water saturation, gas saturation. '
         'Resolution-independent — trained at one grid size, applicable to others.'),
        ('⚖️  PINO Physics Loss',
         'Darcy PDE residuals + Peacemann well model loss + initial/boundary conditions. '
         'Enforces physics between training data points — not just interpolation.'),
        ('🎲  VCAE + DDIM Priors',
         'Variational autoencoder compresses permeability K → latent z. '
         'Denoising Diffusion Implicit Model decodes to non-Gaussian K fields.'),
        ('🌲  CCR Well Model',
         'Cluster-Classify-Regress: XGBoost mixture of experts for the '
         'Peacemann well productivity surrogate. 66 outputs (WOPR/WWPR/WGPR × 22 wells).'),
    ])
]
tech_right = [
    ('History Matching Phase — JAX / XLA (3–5× faster)', GREEN, [
        ('⚡  jax.jit Compiled Kalman Update',
         'Entire Kalman step compiled to XLA kernel. Vectorised over 200 ensemble '
         'members simultaneously. 3–5× faster than equivalent PyTorch loop.'),
        ('🗺️  Gaspari-Cohn Localisation',
         '5th-order piecewise rational function eliminates spurious long-range '
         'correlations in the Kalman gain matrix. Configurable radius.'),
        ('🔄  SVD-Based Matrix Inversion',
         'Numerically stable Kalman gain computation. Truncated singular values '
         'handle near-singular covariance matrices robustly.'),
        ('🎯  Adaptive α (Discrepancy Principle)',
         'Automatic step-size tuning. Guarantees convergence. '
         'Stops when s_cumulative ≥ 1.0 — no manual tuning required.'),
    ])
]

for j, side in enumerate([tech_left, tech_right]):
    c = tech.cell(0, j)
    shade_cell(c, 'F8FAFC' if j == 0 else 'F0FFF4')
    title, col, items = side[0]
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = col; r.font.name = 'Calibri'
    for name, desc in items:
        pn = c.add_paragraph()
        pn.paragraph_format.space_before = Pt(4)
        pn.paragraph_format.space_after  = Pt(2)
        rn = pn.add_run(f'{name}')
        rn.font.bold = True; rn.font.size = Pt(9.5); rn.font.color.rgb = DARK_GREY; rn.font.name = 'Calibri'
        pd = c.add_paragraph(desc)
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after  = Pt(6)
        pd.paragraph_format.left_indent  = Cm(0.4)
        pd.runs[0].font.size = Pt(9); pd.runs[0].font.color.rgb = MID_GREY; pd.runs[0].font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 5: MARKET OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 5, 16, '📈  MARKET OPPORTUNITY', AMBER)
add_heading(doc, 'A $12 Billion Market Ready for Disruption', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Reservoir simulation software is a mature, high-margin, under-innovated market.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=10)

mkt = make_table(doc, 1, 2,
    col_widths=[CONTENT_WIDTH // 2 - Cm(0.3), CONTENT_WIDTH // 2 - Cm(0.3)])

# Left: TAM/SAM/SOM
cl = mkt.cell(0, 0)
shade_cell(cl, 'EBF5FB')
tdata = [
    ('$12B+', 'TAM — Global Reservoir Simulation Software + Services', BLUE, 'D6EAF8'),
    ('$3.5B', 'SAM — History Matching Workflow Automation Tools', CYAN, 'D1F2EB'),
    ('$420M', 'SOM — Year 5 Revenue Target (conservative 1% penetration)', GREEN, 'EAFAF1'),
]
p = cl.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Market Size (TAM / SAM / SOM)')
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE; r.font.name = 'Calibri'

for num, label, col, _ in tdata:
    p2 = cl.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    r2 = p2.add_run(f'{num}   ')
    r2.font.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = col; r2.font.name = 'Calibri'
    r3 = p2.add_run(label)
    r3.font.size = Pt(9); r3.font.color.rgb = DARK_GREY; r3.font.name = 'Calibri'

# Right: market drivers
cr = mkt.cell(0, 1)
shade_cell(cr, 'F8F9FA')
drivers = [
    ('~7,000', 'Active reservoir simulation licenses globally (2024)', CYAN),
    ('8.2% CAGR', 'Market growth rate 2024–2030 (AI-driven acceleration)', GREEN),
    ('$180B+', 'E&P digital transformation investment by 2030', AMBER),
]
p0 = cr.add_paragraph()
p0.paragraph_format.space_before = Pt(6)
r0 = p0.add_run('Key Market Indicators')
r0.font.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = NAVY; r0.font.name = 'Calibri'

for num, label, col in drivers:
    p2 = cr.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(num)
    r2.font.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = col; r2.font.name = 'Calibri'
    p3 = cr.add_paragraph(label)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(6)
    p3.runs[0].font.size = Pt(9); p3.runs[0].font.color.rgb = DARK_GREY; p3.runs[0].font.name = 'Calibri'

p_drivers = cr.add_paragraph()
p_drivers.paragraph_format.space_before = Pt(8)
r_d = p_drivers.add_run('Key Drivers:  ')
r_d.font.bold = True; r_d.font.size = Pt(9.5); r_d.font.color.rgb = AMBER; r_d.font.name = 'Calibri'
r_d2 = p_drivers.add_run('Energy transition pressure → maximise recovery from existing fields. '
    'AI adoption mandates from NOCs. ESG reporting requires rigorous uncertainty quantification.')
r_d2.font.size = Pt(9); r_d2.font.color.rgb = DARK_GREY; r_d2.font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 6: COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 6, 16, '⚔️  COMPETITIVE LANDSCAPE', CYAN)
add_heading(doc, 'We Win Where It Counts Most', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Speed, uncertainty quantification, and AI assistance — three gaps no competitor has filled.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

headers = ['Capability', 'PhysicsFlow ⚡', 'Eclipse 100\n(SLB)',
           'tNavigator\n(RFD)', 'CMG IMEX', 'REVEAL\n(Petex)', 'OPM FLOW\n(Open Source)']
col_w = [Cm(4.8), Cm(3.6), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
comp = make_table(doc, 9, 7, col_widths=col_w)

rows_data = [
    ('HM simulation speed',
     '~7 seconds',
     '3–8 hours', '1–4 hours', '2–6 hours', '1–3 hours', '2–8 hours'),
    ('Ensemble UQ (P10/P50/P90)',
     '200 members, ~1hr',
     'Manual / expensive', 'Limited', 'Add-on only', 'Basic', 'None'),
    ('Built-in AI assistant',
     'Yes — local LLM',
     'No', 'No', 'No', 'No', 'No'),
    ('HPC infrastructure needed',
     'Laptop GPU only',
     'HPC cluster', 'HPC cluster', 'HPC cluster', 'Workstation', 'HPC cluster'),
    ('Annual license cost',
     '$15k – $50k',
     '$150k – $400k', '$50k – $200k', '$100k – $300k', '$80k – $200k', 'Free'),
    ('Physics compliance',
     'PDE loss enforced', 'Full FVM', 'Full FVM', 'Full FVM', 'Full FVM', 'Full FVM'),
    ('Generative priors (VCAE/DDIM)',
     'Yes', 'No', 'No', 'No', 'No', 'No'),
    ('Eclipse format support',
     'Native import', 'Native', 'Full', 'Partial', 'Partial', 'Native'),
]

# Header row
for j, h in enumerate(headers):
    shade_cell(comp.cell(0, j), '0A1628' if j != 1 else '0D3B6E')
    p = comp.cell(0, j).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.bold = True; run.font.size = Pt(8.5)
    run.font.color.rgb = CYAN if j == 1 else WHITE
    run.font.name = 'Calibri'

# Data rows
for i, row in enumerate(rows_data):
    bg_even = 'F8FAFC'
    for j, val in enumerate(row):
        is_pf = (j == 1)
        bg = 'EBF5FB' if is_pf else (bg_even if i % 2 == 0 else 'FFFFFF')
        shade_cell(comp.cell(i+1, j), bg)
        cell_vertical_align(comp.cell(i+1, j), 'center')
        p = comp.cell(i+1, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(9 if j == 0 else 8.5)
        run.font.bold = is_pf
        run.font.name = 'Calibri'
        if is_pf:
            run.font.color.rgb = BLUE
        elif j == 0:
            run.font.color.rgb = DARK_GREY
        else:
            run.font.color.rgb = MID_GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 7: KEY ADVANTAGES
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 7, 16, '🏆  COMPETITIVE ADVANTAGES', CYAN)
add_heading(doc, 'Four Moats That Are Hard to Copy', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Incumbents cannot easily retrofit AI into 40-year-old finite-volume codebases.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

adv = make_table(doc, 2, 2,
    col_widths=[CONTENT_WIDTH//2 - Cm(0.3), CONTENT_WIDTH//2 - Cm(0.3)])
adv_data = [
    ('⚡  6,000× Speed Moat — Physics-Informed Architecture', 'EBF5FB', BLUE,
     'The PINO surrogate is not a black box — it is trained with Darcy PDE residuals as a loss term. '
     'This is fundamentally different from ML curve-fitting. Incumbent vendors cannot add this to '
     'their FVM solvers without a ground-up rewrite. Our speed advantage compounds: faster runs → '
     'more ensemble members → better UQ → better decisions → more operator value.'),
    ('📦  Data Moat — Growing Training Library', 'F0FFF4', GREEN,
     'Every new field a customer brings generates additional OPM FLOW training data. Over time, '
     'PhysicsFlow builds a proprietary library of pre-trained surrogates for common geological settings '
     '— reducing training time from 5 hours to minutes via transfer learning. '
     'This flywheel is impossible for new entrants to replicate quickly.'),
    ('🤖  AI Assistant Moat — Embedded Knowledge', 'FFFBF0', AMBER,
     'No existing reservoir simulation tool has a natural-language interface grounded in live '
     'simulation data. Engineers who learn to ask questions in plain English will not want to go back '
     'to writing scripts and digging through spreadsheets. The AI assistant becomes a productivity '
     'multiplier that creates deep product stickiness and reduces support burden.'),
    ('💰  Economics Moat — 10× Cost Advantage', 'FFF8F8', RED,
     'At $15,000/year Professional tier vs $150,000+ for Eclipse, PhysicsFlow opens the market to '
     '10× more operators — particularly independents and NOCs in developing markets. '
     'Once a customer builds their workflow around PhysicsFlow, switching costs are high '
     '(trained models, project files, learned workflows). Low entry price drives high retention.'),
]
coords = [(0,0), (0,1), (1,0), (1,1)]
for (ri, ci), (title, bg, col, desc) in zip(coords, adv_data):
    c = adv.cell(ri, ci)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = col; r.font.name = 'Calibri'
    p2 = c.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(8)
    p2.runs[0].font.size = Pt(9.5); p2.runs[0].font.color.rgb = DARK_GREY; p2.runs[0].font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_body(doc,
    'Scientific foundation: Built on peer-reviewed NVIDIA research (arXiv:2406.00889, 2024), '
    'validated on the real Norne Field benchmark — a publicly available Norwegian North Sea dataset. '
    'This is not vaporware. The physics and algorithms are published, peer-reviewed, and reproducible.',
    color=BLUE, size=9.5, bold=False, italic=True, space_before=0, space_after=0)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 8: TARGET CUSTOMERS
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 8, 16, '🎯  TARGET CUSTOMERS', AMBER)
add_heading(doc, 'Three High-Value Segments', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Prioritised by ease of sale, willingness to pay, and strategic value.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

cust = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2)])
cust_data = [
    ('🏭  Independent Operators\n& E&P Consultancies', 'EBF5FB', BLUE, 'Priority 1 — Immediate',
     'Companies producing 1,000–50,000 BOE/day. Cannot afford $200,000+ Eclipse licenses or HPC infrastructure. PhysicsFlow gives them enterprise-grade history matching at 10% of the cost.',
     [('~3,500 companies globally', BLUE),
      ('Average deal: $15,000–$25,000/yr', BLUE),
      ('Sales cycle: 1–3 months', BLUE),
      ('Key pain: cost + speed', BLUE)]),
    ('🏛️  National Oil Companies\n(NOCs)', 'F0FFF4', GREEN, 'Priority 2 — 12–18 months',
     'ADNOC, Petronas, ONGC, ENAP, Pertamina — all under mandate to digitise and apply AI. PhysicsFlow\'s speed and AI assistant align perfectly with their digital transformation programmes.',
     [('~60 major NOCs worldwide', GREEN),
      ('Average deal: $100k–$500k/yr', GREEN),
      ('Sales cycle: 6–18 months', GREEN),
      ('Key pain: UQ + AI adoption', GREEN)]),
    ('🔬  Service Companies\n& Engineering Consultants', 'FFFBF0', AMBER, 'Priority 3 — 18–36 months',
     'Halliburton, Baker Hughes, Gaffney Cline, Ryder Scott — firms that run HM studies for clients. 10× throughput increase means more projects, same staff, dramatically higher margins.',
     [('~800 firms globally', AMBER),
      ('Average deal: $50k–$200k/yr', AMBER),
      ('Sales cycle: 3–9 months', AMBER),
      ('OEM/white-label opportunity', AMBER)]),
]

for j, (title, bg, col, priority, desc, metrics) in enumerate(cust_data):
    c = cust.cell(0, j)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    pp = c.add_paragraph()
    pp.paragraph_format.space_before = Pt(6)
    r_pri = pp.add_run(f'{priority}\n')
    r_pri.font.size = Pt(8); r_pri.font.bold = True; r_pri.font.color.rgb = col; r_pri.font.name = 'Calibri'
    r_tit = c.add_paragraph()
    rt = r_tit.add_run(title)
    rt.font.bold = True; rt.font.size = Pt(12); rt.font.color.rgb = NAVY; rt.font.name = 'Calibri'
    r_tit.paragraph_format.space_after = Pt(6)
    pd = c.add_paragraph(desc)
    pd.paragraph_format.space_after = Pt(8)
    pd.runs[0].font.size = Pt(9.5); pd.runs[0].font.color.rgb = DARK_GREY; pd.runs[0].font.name = 'Calibri'
    for m_text, m_col in metrics:
        pm = c.add_paragraph()
        pm.paragraph_format.space_before = Pt(2)
        pm.paragraph_format.space_after  = Pt(2)
        rm = pm.add_run(f'◆  {m_text}')
        rm.font.size = Pt(9); rm.font.color.rgb = m_col; rm.font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 9: BUSINESS MODEL
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 9, 16, '💼  BUSINESS MODEL', CYAN)
add_heading(doc, 'SaaS + Perpetual · Multiple Revenue Streams', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Annual subscription below competitors with high expansion revenue from cloud compute and professional services.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

price = make_table(doc, 6, 4,
    col_widths=[CONTENT_WIDTH//4 - Cm(0.2), CONTENT_WIDTH//4 - Cm(0.2),
                CONTENT_WIDTH//4 - Cm(0.2), CONTENT_WIDTH//4 - Cm(0.2)])

price_data = [
    ('COMMUNITY', 'Free', 'forever', 'D6EAF8', MID_GREY,
     ['1 field / project', '50 ensemble members', 'Basic PINO (FNO mode)', 'No CCR well model', 'Community support', 'Target: Universities, students']),
    ('PROFESSIONAL  ★', '$15,000', 'per year · 1 seat', 'D5F5E3', GREEN,
     ['Unlimited fields', '500 ensemble members', 'Full PINO + CCR + αREKI', 'VCAE + DDIM priors', 'AI assistant (Ollama)', 'Eclipse import/export', 'Email support (48hr SLA)']),
    ('ENTERPRISE', '$50,000', 'per year · 5 seats', 'D6EAF8', BLUE,
     ['Everything in Professional', '1,000 ensemble members', 'Multi-user project sharing', 'REST API access', 'VTK export (ResInsight)', 'Dedicated support + SLA', 'Custom training workshops']),
    ('CLOUD / PAY-PER-USE', '$5 – $20', 'per simulation run', 'FFF9E6', AMBER,
     ['No installation required', 'Fully managed GPU cloud', 'Scale to 2,000 members', 'Results streamed to browser', 'Pay only for compute used']),
]

# Header row
for j, (tier, amount, period, bg, col, _) in enumerate(price_data):
    shade_cell(price.cell(0, j), bg)
    shade_cell(price.cell(1, j), bg)
    shade_cell(price.cell(2, j), bg)
    cell_para(price.cell(0, j), tier, bold=True, size=9, color=col, center=True)
    cell_para(price.cell(1, j), amount, bold=True, size=24, color=NAVY, center=True)
    cell_para(price.cell(2, j), period, bold=False, size=8, color=MID_GREY, center=True)

# Features
for i in range(3, 6):
    for j in range(4):
        shade_cell(price.cell(i, j), 'FFFFFF' if i % 2 == 0 else 'F8FAFC')

for j, (_, _, _, bg, col, features) in enumerate(price_data):
    for fi, feat in enumerate(features[:4]):
        if 3 + fi < 7:
            shade_cell(price.cell(3 + (fi // 2), j), 'FFFFFF' if fi % 2 == 0 else 'F8FAFC')
    c = price.cell(3, j)
    c.merge(price.cell(5, j))
    shade_cell(c, 'FAFAFA')
    for feat in features:
        pf = c.add_paragraph(f'✓  {feat}')
        pf.paragraph_format.space_before = Pt(2)
        pf.paragraph_format.space_after  = Pt(2)
        pf.runs[0].font.size = Pt(8.5); pf.runs[0].font.color.rgb = DARK_GREY; pf.runs[0].font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Additional revenue streams
add_body(doc, 'Additional Revenue Streams:', color=NAVY, size=10, bold=True, space_before=0, space_after=4)
rev = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2)])
rev_data = [
    ('💰  Professional Services', 'D6EAF8', BLUE,
     'Field-specific PINO training, custom integrations, on-site workshops. $2,500–$10,000/day consulting.'),
    ('🤝  OEM / White-Label', 'D5F5E3', GREEN,
     'License the engine to service companies embedded in their workflows. Revenue share model.'),
    ('📚  Pre-Trained Surrogates', 'FFF9E6', AMBER,
     'Marketplace of field-type surrogates (carbonate, clastic, fractured). $5k–$20k per surrogate.'),
]
for j, (title, bg, col, desc) in enumerate(rev_data):
    shade_cell(rev.cell(0, j), bg)
    p = rev.cell(0, j).add_paragraph()
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = col; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    pd = rev.cell(0, j).add_paragraph(desc)
    pd.runs[0].font.size = Pt(9); pd.runs[0].font.color.rgb = DARK_GREY; pd.runs[0].font.name = 'Calibri'
    pd.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 10: FINANCIAL PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 10, 16, '📊  FINANCIAL PROJECTIONS', GREEN)
add_heading(doc, 'Path to $42M ARR by Year 5', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Conservative model based on 1% penetration of addressable market — realistic with direct sales plus channel partners.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

fin = make_table(doc, 9, 6,
    col_widths=[Cm(6.0), Cm(3.5), Cm(3.5), Cm(3.5), Cm(4.0), Cm(4.0)])

fin_headers = ['Metric', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5']
fin_rows = [
    ('Professional customers',      '15',       '60',       '180',      '420',      '800'),
    ('Enterprise customers',        '2',        '8',        '22',       '50',       '120'),
    ('Cloud run revenue',           '$50k',     '$200k',    '$800k',    '$2.5M',    '$6M'),
    ('Total ARR',                   '$425k',    '$2.1M',    '$6.4M',    '$16.2M',   '$42M'),
    ('Gross margin',                '72%',      '78%',      '82%',      '84%',      '86%'),
    ('EBITDA',                      '($1.8M)',  '($1.2M)',  '$0.4M',    '$3.8M',    '$14.7M'),
    ('Headcount',                   '6',        '14',       '28',       '45',       '70'),
    ('Cumulative investment needed', '$3.5M',   '$3.5M',    'Self-fund', 'Series A', 'Series A'),
]
highlight_rows = {3, 5}

for j, h in enumerate(fin_headers):
    shade_cell(fin.cell(0, j), '0A1628')
    p = fin.cell(0, j).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(h)
    run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = CYAN; run.font.name = 'Calibri'

for i, (row) in enumerate(fin_rows):
    is_hi = i in highlight_rows
    for j, val in enumerate(row):
        bg = 'EBF5FB' if is_hi else ('F8FAFC' if i % 2 == 0 else 'FFFFFF')
        shade_cell(fin.cell(i+1, j), bg)
        p = fin.cell(i+1, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(9); run.font.bold = is_hi and j > 0; run.font.name = 'Calibri'
        is_negative = val.startswith('(')
        is_positive = (is_hi and j in (4, 5)) or (i == 5 and j in (3, 4, 5) and not is_negative)
        if is_negative: run.font.color.rgb = RED
        elif is_positive: run.font.color.rgb = GREEN
        elif is_hi: run.font.color.rgb = BLUE
        else: run.font.color.rgb = DARK_GREY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Key metrics
km = make_table(doc, 1, 4,
    col_widths=[CONTENT_WIDTH//4]*4)
km_data = [
    ('86%', 'Target gross margin (Year 5)', CYAN, 'D6EAF8'),
    ('Year 3', 'EBITDA breakeven', GREEN, 'D5F5E3'),
    ('$600k', 'Average enterprise ACV', AMBER, 'FFF9E6'),
    ('12×', 'Target EV/ARR multiple (Year 5)', BLUE, 'EBF5FB'),
]
for j, (num, label, col, bg) in enumerate(km_data):
    shade_cell(km.cell(0, j), bg)
    cell_vertical_align(km.cell(0, j), 'center')
    p = km.cell(0, j).add_paragraph(num)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True; p.runs[0].font.size = Pt(24); p.runs[0].font.color.rgb = col; p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    p2 = km.cell(0, j).add_paragraph(label)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(8); p2.runs[0].font.color.rgb = MID_GREY; p2.runs[0].font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 11: GO-TO-MARKET
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 11, 16, '🚀  GO-TO-MARKET STRATEGY', AMBER)
add_heading(doc, 'Land, Expand, Dominate', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Start with independents (fast sales, quick wins), expand to NOCs and service companies via channel.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

gtm = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2), CONTENT_WIDTH//3 - Cm(0.2)])
gtm_data = [
    ('Phase 1  ·  Months 1–12', 'Land: Quick Wins', 'D6EAF8', BLUE, [
        'Free Norne benchmark demo — every prospect runs it Day 1',
        'Target 15 Professional accounts via direct outreach at SPE/ADIPEC',
        'Publish 3 peer-reviewed case studies with early adopters',
        'GitHub open-source community tier — build developer mindshare',
        'University partnerships (Heriot-Watt, Imperial, KFUPM)',
        'Free webinar series: "History matching in 1 hour" — lead gen',
    ]),
    ('Phase 2  ·  Months 12–24', 'Expand: Channel & NOCs', 'D5F5E3', GREEN, [
        'Hire 4 technical sales engineers (ex-reservoir engineers)',
        'Reseller agreements: Middle East, SE Asia, LatAm consultancies',
        'First NOC pilot — ADNOC or Petronas technology partnership MoU',
        'Integration with ResInsight — credibility with Equinor ecosystem',
        'Pre-trained surrogate marketplace launch',
        'Cloud product launch (pay-per-run, no install barrier)',
    ]),
    ('Phase 3  ·  Months 24–48', 'Dominate: Enterprise', 'FFF9E6', AMBER, [
        'Full enterprise sales — dedicated AEs per region',
        'OEM licensing deal with 1–2 major oilfield service companies',
        'ISO 9001 quality certification for regulatory acceptance',
        'SPE reserves certification pathway — enables SEC filings',
        'Compositional + geomechanics extension: deepwater/unconventional',
        'Cloud platform becomes majority of revenue',
    ]),
]

for j, (phase, title, bg, col, items) in enumerate(gtm_data):
    c = gtm.cell(0, j)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    pp = c.add_paragraph()
    pp.paragraph_format.space_before = Pt(6)
    r_ph = pp.add_run(phase)
    r_ph.font.size = Pt(8); r_ph.font.bold = True; r_ph.font.color.rgb = col; r_ph.font.name = 'Calibri'
    pt = c.add_paragraph()
    rt = pt.add_run(title)
    rt.font.bold = True; rt.font.size = Pt(13); rt.font.color.rgb = NAVY; rt.font.name = 'Calibri'
    pt.paragraph_format.space_after = Pt(8)
    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.space_before = Pt(2)
        pi.paragraph_format.space_after  = Pt(2)
        ri = pi.add_run(f'◆  {item}')
        ri.font.size = Pt(9); ri.font.color.rgb = DARK_GREY; ri.font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 12: PRODUCT ROADMAP
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 12, 16, '🗺️  PRODUCT ROADMAP', CYAN)
add_heading(doc, '18-Month Build Plan', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'From validated research prototype to full commercial product ready for the oil and gas market.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

rm = make_table(doc, 1, 4,
    col_widths=[CONTENT_WIDTH//4 - Cm(0.2)] * 4)
rm_data = [
    ('✅  Now — Foundation', CYAN, 'D6EAF8', [
        'Python gRPC engine', 'JAX αREKI (3–5× faster)', 'Ollama AI assistant',
        '.NET 8 WPF desktop shell', 'Protocol Buffer API contracts',
        'PVT + well model (config-driven)', 'Dashboard + HM + Training views', 'AI chat panel',
    ]),
    ('🔄  Q2 2025 — v1.1 Complete', GREEN, 'D5F5E3', [
        'FNO surrogate (PyTorch)', 'PINO training service', 'Eclipse deck reader (.DATA/.EGRID)',
        'VCAE + DDIM priors', 'CCR well model (XGBoost)', '3D field viewer (Helix Toolkit)',
        'PDF / Excel report export', 'WiX professional installer',
    ]),
    ('🔮  Q4 2025 — v1.2 Production', AMBER, 'FFF9E6', [
        'Compositional PVT (EOS)', 'Horizontal / multilateral wells', 'Surface network (lite)',
        'Multi-user project sharing', 'REST API for integrations', 'Cloud deployment (SaaS)',
        'Audit trail — ISO 9001 path', 'Auto-update delivery system',
    ]),
    ('🚀  2026 — v2.0 Platform', BLUE, 'EBF5FB', [
        'Geomechanics coupling', 'Thermal / EOR simulation', 'Eclipse 100/300 bridge API',
        'Pre-trained surrogate marketplace', 'Transfer learning (new fields in <30min)',
        'Multi-GPU distributed training', 'SPE reserve certification ready', 'Petrel plugin',
    ]),
]

for j, (title, col, bg, items) in enumerate(rm_data):
    c = rm.cell(0, j)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = col; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(8)
    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.space_before = Pt(2)
        pi.paragraph_format.space_after  = Pt(2)
        ri = pi.add_run(f'→  {item}')
        ri.font.size = Pt(9); ri.font.color.rgb = DARK_GREY; ri.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_body(doc,
    'Build risk is LOW: the hard research problems are solved (published, peer-reviewed). '
    'Remaining work is engineering — UI, I/O, installer, testing — well-understood and estimable.',
    color=GREEN, size=9.5, italic=True, bold=False, space_before=0)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 13: TEAM
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 13, 16, '👥  THE TEAM', CYAN)
add_heading(doc, 'Built by Reservoir Engineers and AI Scientists', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'The rare team that understands both the physics and the software — an essential combination in petroleum AI.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

team = make_table(doc, 1, 4,
    col_widths=[CONTENT_WIDTH//4 - Cm(0.2)] * 4)
team_data = [
    ('CEO / Reservoir Engineering Lead', 'Founder', 'D6EAF8', BLUE,
     '15+ years reservoir simulation. Former senior engineer at major operator. MSc Petroleum Engineering. Deep expertise in history matching and UQ. Relationships with key NOC buyers.'),
    ('CTO / ML Architecture Lead', 'Co-Founder', 'D5F5E3', GREEN,
     'PhD Computational Physics. Expert in physics-informed ML, neural operators, JAX/PyTorch. Former NVIDIA research team. Author of published FNO and PINO extensions.'),
    ('Head of Product / Platform', 'To Hire — Q1 2025', 'FFF9E6', AMBER,
     '10+ years engineering software development. WPF/.NET expertise. Experience building industrial desktop applications for the oil and gas sector. UX-first mindset.'),
    ('VP Sales — Oil & Gas', 'To Hire — Q2 2025', 'EBF5FB', MID_GREY,
     'Former technical sales at SLB or Petex. Existing relationships with NOC and independent operator procurement teams. MENA + Asia Pacific regional focus.'),
]

for j, (name, role, bg, col, bio) in enumerate(team_data):
    c = team.cell(0, j)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    p1 = c.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    r1 = p1.add_run(name)
    r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NAVY; r1.font.name = 'Calibri'
    p2 = c.add_paragraph()
    r2 = p2.add_run(role)
    r2.font.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = col; r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(6)
    p3 = c.add_paragraph(bio)
    p3.runs[0].font.size = Pt(9); p3.runs[0].font.color.rgb = DARK_GREY; p3.runs[0].font.name = 'Calibri'
    p3.paragraph_format.space_after = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

partners = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2)] * 3)
partner_data = [
    ('🏛️  Scientific Advisory Board', 'D6EAF8', BLUE,
     'Recruiting 3 advisors: (1) SPE Distinguished Member — HM expertise. (2) Professor of Computational Geoscience. (3) Former CTO of major reservoir software company.'),
    ('🤝  Strategic Partners', 'D5F5E3', GREEN,
     'OPM Project · NVIDIA PhysicsNeMo team · Equinor Norne benchmark committee · Microsoft for Startups (Azure GPU credits) · SPE Digital Energy Technical Section'),
    ('📜  IP Position', 'FFF9E6', AMBER,
     'Core algorithms are published open science — competitive moat is implementation, data, UX, and domain expertise. Provisional patent pending on JAX αREKI acceleration.'),
]
for j, (title, bg, col, desc) in enumerate(partner_data):
    c = partners.cell(0, j)
    shade_cell(c, bg)
    p1 = c.add_paragraph()
    p1.paragraph_format.space_before = Pt(6)
    r1 = p1.add_run(title)
    r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = col; r1.font.name = 'Calibri'
    p2 = c.add_paragraph(desc)
    p2.runs[0].font.size = Pt(9); p2.runs[0].font.color.rgb = DARK_GREY; p2.runs[0].font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 14: RISK ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 14, 16, '⚠  RISK ANALYSIS', RED)
add_heading(doc, 'Risks We Have Thought About — And How We Mitigate Them', color=NAVY, size=28, space_before=0, space_after=4)
add_body(doc, 'Honest assessment: the risks are real but manageable. Technology risk is essentially zero.',
         color=DARK_GREY, size=11, italic=True, space_before=0, space_after=8)

risks = make_table(doc, 5, 3,
    col_widths=[Cm(5.5), Cm(6.0), Cm(6.0)])

risk_headers = ['Risk', 'Description', 'Mitigation']
for j, h in enumerate(risk_headers):
    shade_cell(risks.cell(0, j), '0A1628')
    p = risks.cell(0, j).paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = CYAN; r.font.name = 'Calibri'

risk_data = [
    ('Physics accuracy vs. full FVM',
     'PINO surrogate has approximation error compared to full Eclipse/OPM FLOW physics — a legitimate concern for reserve certification.',
     'Surrogate error quantified and propagated into αREKI. Positioned as fast-screening tool; full FVM recommended for final reserve booking. SPE endorsement roadmap.'),
    ('SLB / incumbent counter-move',
     'SLB has capital to develop competing AI HM tools. They are aware of PINO literature and have internal AI research teams.',
     '18–24 month head start before SLB can ship a comparable product. Price advantage ($15k vs $150k) protects downmarket. Data flywheel creates durable network effects.'),
    ('Regulatory acceptance',
     'SEC and SPE-PRMS reserve certification requires defensible methodology. Regulators are conservative and unfamiliar with neural operator surrogates.',
     'Full audit trail, reproducibility (seed stored in project file), SVD-based numerical stability. Certify for screening first, expand to booking as acceptance grows.'),
    ('Training data requirements',
     '100 OPM FLOW runs required per field — some operators may resist running a simulator before seeing value from the surrogate.',
     'Pre-trained surrogates for common geological settings. Cloud burst GPU option. ROI argument: 100 training runs × 7 sec inference = full HM in 1 hour vs 6 months manually.'),
]

for i, (risk, desc, mitigation) in enumerate(risk_data):
    bg = 'FFF8F8' if i % 2 == 0 else 'FFFFFF'
    shade_cell(risks.cell(i+1, 0), 'FFF3F3')
    shade_cell(risks.cell(i+1, 1), bg)
    shade_cell(risks.cell(i+1, 2), 'F0FFF4')
    cell_vertical_align(risks.cell(i+1, 0), 'top')
    cell_vertical_align(risks.cell(i+1, 1), 'top')
    cell_vertical_align(risks.cell(i+1, 2), 'top')
    p0 = risks.cell(i+1, 0).paragraphs[0]
    r0 = p0.add_run(f'⚠  {risk}')
    r0.font.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = RED; r0.font.name = 'Calibri'
    p1 = risks.cell(i+1, 1).paragraphs[0]
    r1 = p1.add_run(desc)
    r1.font.size = Pt(9); r1.font.color.rgb = DARK_GREY; r1.font.name = 'Calibri'
    p2 = risks.cell(i+1, 2).paragraphs[0]
    r2 = p2.add_run(f'✓  {mitigation}')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20); r2.font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 15: THE ASK
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 15, 16, '💰  THE ASK', GREEN)
add_heading(doc, '$3.5M Seed Round', color=NAVY, size=40, space_before=0, space_after=6, center=True)
add_body(doc, '18 months runway  ·  Product completion  ·  First 30 customers  ·  Series A ready at $2.1M ARR',
         color=MID_GREY, size=12, italic=True, space_before=0, space_after=14, center=True)

use = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2)] * 3)
use_data = [
    ('40% — Engineering & Product\n$1.4M', BLUE, 'D6EAF8',
     '3 senior engineers covering: FNO training service, Eclipse I/O, 3D reservoir viewer, WiX installer, cloud deployment, test suite, and documentation. 18 months full product build.'),
    ('30% — Sales & Marketing\n$1.05M', GREEN, 'D5F5E3',
     '2 technical sales engineers (former reservoir engineers). SPE/ADIPEC conference presence. Case study production. Demand generation. First NOC pilot engagement budget.'),
    ('20% — Compute & Infrastructure\n$700k', AMBER, 'FFF9E6',
     'GPU cloud for customer training runs. Pre-trained surrogate library generation. CI/CD infrastructure build-out. Third-party security audit. 10% reserved for contingency.'),
]
for j, (title, col, bg, desc) in enumerate(use_data):
    c = use.cell(0, j)
    shade_cell(c, bg)
    cell_vertical_align(c, 'top')
    p1 = c.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    r1 = p1.add_run(title)
    r1.font.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = col; r1.font.name = 'Calibri'
    p2 = c.add_paragraph(desc)
    p2.runs[0].font.size = Pt(9.5); p2.runs[0].font.color.rgb = DARK_GREY; p2.runs[0].font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Milestones
milestones = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2)] * 3)
ms_data = [
    ('30 Customers', 'Target by Month 18', CYAN, 'D6EAF8'),
    ('$2.1M ARR', 'Series A Milestone', GREEN, 'D5F5E3'),
    ('v1.2 Shipped', 'Full Production-Ready Product', AMBER, 'FFF9E6'),
]
for j, (num, label, col, bg) in enumerate(ms_data):
    c = milestones.cell(0, j)
    shade_cell(c, bg)
    p1 = c.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(8)
    r1 = p1.add_run(num)
    r1.font.bold = True; r1.font.size = Pt(26); r1.font.color.rgb = col; r1.font.name = 'Calibri'
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(label)
    r2.font.size = Pt(9); r2.font.color.rgb = MID_GREY; r2.font.name = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 16: CLOSING
# ══════════════════════════════════════════════════════════════════════════════

add_slide_header(doc, 16, 16, '🌍  LET\'S BUILD IT TOGETHER', CYAN)
add_heading(doc, 'The Vision', color=NAVY, size=28, space_before=0, space_after=10, center=True)

# Big quote
q = make_table(doc, 1, 1, col_widths=[CONTENT_WIDTH])
shade_cell(q.cell(0, 0), 'EBF5FB')
c = q.cell(0, 0)
p = c.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('"The oil and gas industry produces $3 trillion of value annually.\n'
              'Improving reservoir recovery by even 1% is worth $30 billion.\n'
              'PhysicsFlow makes that 1% improvement systematic, fast,\n'
              'and accessible to every operator — not just the supermajors."')
r.font.size = Pt(16); r.font.italic = True; r.font.color.rgb = NAVY; r.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(20)

doc.add_paragraph().paragraph_format.space_after = Pt(16)

add_divider(doc, color='1565C0')

add_heading(doc, 'Let\'s build the future of reservoir engineering together.',
            color=NAVY, size=20, bold=True, space_before=12, space_after=16, center=True)

contact = make_table(doc, 1, 3,
    col_widths=[CONTENT_WIDTH//3 - Cm(0.2)] * 3)
contact_data = [
    ('📧  Email', 'contact@physicsflow.ai', BLUE, 'D6EAF8'),
    ('🌐  Website', 'physicsflow.ai', GREEN, 'D5F5E3'),
    ('💻  GitHub', 'github.com/physicsflow', CYAN, 'EBF5FB'),
]
for j, (label, val, col, bg) in enumerate(contact_data):
    c = contact.cell(0, j)
    shade_cell(c, bg)
    p1 = c.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(10)
    r1 = p1.add_run(label)
    r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = col; r1.font.name = 'Calibri'
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(val)
    r2.font.size = Pt(11); r2.font.color.rgb = NAVY; r2.font.name = 'Calibri'


# ── Save ──────────────────────────────────────────────────────────────────────
output = r'C:\Users\danie\OneDrive\Reservoir_Simulator\PhysicsFlow\PhysicsFlow_PitchDeck.docx'
doc.save(output)
print(f'Saved: {output}')
